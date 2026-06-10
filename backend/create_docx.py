from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# Set up styles
style = doc.styles['Heading 1']
style.font.size = Pt(24)
style.font.bold = True
style.font.color.rgb = RGBColor(37, 99, 235)

style2 = doc.styles['Heading 2']
style2.font.size = Pt(18)
style2.font.bold = True
style2.font.color.rgb = RGBColor(15, 45, 74)

# Title Page
title = doc.add_heading('AutiStudy', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('AI-Powered Adaptive Learning Platform for Students with Autism')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
p = doc.add_paragraph('Personalized tutoring for Grades 4-7 in Pakistan')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('Aligned with Pakistan\'s National Curriculum')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('Bilingual Support: English & Urdu')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# Slide 2
doc.add_heading('Slide 2: Live Demonstration', 1)
doc.add_heading('Live Demo of AutiStudy', 2)
items = [
    'User registration and login flow',
    'Dashboard overview with navigation',
    'Subject and grade selection',
    'AI Tutor chat interaction',
    'Auto-generated visual explanations',
    'Text-to-Speech functionality',
    'Practice quizzes with analytics',
    'Bilingual interface switching'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph('[Live demonstration of the Streamlit application]', style='Intense Quote')
doc.add_page_break()

# Slide 3
doc.add_heading('Slide 3: Problem Statement', 1)
doc.add_heading('Why Do We Need This?', 2)
items = [
    'Students with autism have unique learning needs that traditional education often fails to address',
    'Lack of patient, adaptive, and personalized tutoring resources in Pakistan',
    'Limited access to quality education for special needs students',
    'No existing platform combines AI tutoring with Pakistan\'s national curriculum',
    'Need for multi-modal learning: text, visuals, and audio together'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_page_break()

# Slide 4
doc.add_heading('Slide 4: Gap Analysis', 1)
doc.add_heading('Current Educational Gaps', 2)
table = doc.add_table(rows=7, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Gap'
hdr[1].text = 'Impact'
data = [
    ('One-size-fits-all teaching', 'Doesn\'t accommodate different learning speeds'),
    ('Complex language in textbooks', 'Confuses students who need simple explanations'),
    ('Lack of visual aids', 'Misses visual learners, especially in Maths'),
    ('No patience in tutoring', 'Rushes students, causes anxiety'),
    ('English-only platforms', 'Excludes Urdu-speaking students'),
    ('Generic AI tutors', 'Not aligned with local curriculum')
]
for i, (gap, impact) in enumerate(data):
    table.rows[i+1].cells[0].text = gap
    table.rows[i+1].cells[1].text = impact
doc.add_page_break()

# Slide 5
doc.add_heading('Slide 5: Our Solution - AutiStudy', 1)
doc.add_heading('What AutiStudy Offers', 2)
items = [
    'AI tutor specifically designed for autism-friendly communication',
    'Step-by-step explanations with infinite patience',
    'Curriculum-aligned content from actual Pakistan textbooks',
    'Auto-generated images for visual understanding',
    'Voice playback for audio learners',
    'Bilingual support (English & Urdu)',
    'Practice quizzes with encouraging feedback',
    'Progress tracking with rewards system'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_page_break()

# Slide 6
doc.add_heading('Slide 6: What is RAG?', 1)
doc.add_heading('Retrieval-Augmented Generation (RAG)', 2)
items = [
    'RAG = Retrieval + Generation combined',
    'Instead of relying only on LLM\'s training data, we retrieve relevant content from textbooks',
    'Ensures answers are accurate and curriculum-based',
    'Prevents hallucination by grounding responses in actual textbook content'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()
doc.add_heading('RAG Pipeline:', 2)
steps = [
    '1. User asks a question',
    '2. System retrieves relevant chunks from textbook database',
    '3. Retrieved context is sent to LLM',
    '4. LLM generates accurate, curriculum-based response'
]
for step in steps:
    doc.add_paragraph(step)
doc.add_page_break()

# Slide 7
doc.add_heading('Slide 7: Initial Approach - Traditional RAG with OCR', 1)
doc.add_heading('First Attempt: OCR-Based Text Extraction', 2)
items = [
    'Used OCR (Optical Character Recognition) to extract text from PDF textbooks',
    'Created basic embeddings from extracted text',
    'Built simple vector database with fixed-size chunks'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()
doc.add_heading('The Reality:', 2)
items = [
    'Simple retrieve → generate pipeline',
    'No understanding of document structure',
    'No awareness of question type or intent',
    'Single retrieval method (only dense search)'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_page_break()

# Slide 8
doc.add_heading('Slide 8: The OCR Disaster - What Went Wrong', 1)
doc.add_heading('Critical Failures in Text Extraction', 2)
table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Content Type'
hdr[1].text = 'What OCR Produced'
hdr[2].text = 'Impact'
data = [
    ('Math: 2/3 + 1/4', '2 3 1 4 or 2/31/4', 'Completely wrong'),
    ('Fractions: ½', '1 2 or 12', 'Lost meaning'),
    ('Symbols: √16 = 4', 'V16 = 4 or √164', 'Garbled'),
    ('Tables', 'Random text jumble', 'Structure lost'),
    ('Urdu mixed', 'Encoding errors', 'Unreadable'),
    ('Diagrams', 'Nothing', 'Ignored')
]
for i, (content, ocr, impact) in enumerate(data):
    table.rows[i+1].cells[0].text = content
    table.rows[i+1].cells[1].text = ocr
    table.rows[i+1].cells[2].text = impact
doc.add_paragraph()
p = doc.add_paragraph('Result: Garbage in = Garbage out')
p.runs[0].bold = True
p = doc.add_paragraph('Retrieval Accuracy: Only 35-45%')
p.runs[0].bold = True
p.runs[0].font.color.rgb = RGBColor(220, 38, 38)
doc.add_page_break()

# Slide 9
doc.add_heading('Slide 9: The Real Problem - Retriever Failure', 1)
doc.add_heading('Why Traditional RAG Failed', 2)
doc.add_paragraph('Student asks: "How to add fractions with different denominators?"')
doc.add_paragraph()
doc.add_paragraph('Traditional RAG Retrieved:')
items = [
    '❌ "The denominator is the bottom number..." (irrelevant definition)',
    '❌ "Fractions are parts of a whole..." (too basic)',
    '❌ "Exercise 2.3: Add 1/2 + garbled..." (OCR corrupted)'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('Expected:')
items = [
    '✅ Step-by-step procedure for finding LCD',
    '✅ Example with worked solution',
    '✅ Practice problems'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()
doc.add_heading('The retriever couldn\'t find the RIGHT content because:', 2)
items = [
    'OCR destroyed mathematical content',
    'No block type awareness (definition vs procedure)',
    'No intent understanding',
    'Single retrieval method missed keywords'
]
for i, item in enumerate(items, 1):
    doc.add_paragraph(f'{i}. {item}')
doc.add_page_break()

# Slide 10
doc.add_heading('Slide 10: The Breakthrough - LlamaParse to Markdown', 1)
doc.add_heading('Research Discovery: Parse, Don\'t OCR', 2)
doc.add_paragraph('After researching alternatives, we discovered LlamaParse (LlamaIndex):')
items = [
    'AI-powered document parsing (not simple OCR)',
    'Understands document structure',
    'Preserves mathematical notation',
    'Outputs clean Markdown files'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()
doc.add_heading('Before (OCR):', 2)
doc.add_paragraph('To add fractions 1 2 + 1 3 first find LCD\nThe LCD of 2 and 3 is 6\nSo 1 2 = 3 6 and 1 3 = 2 6')
doc.add_paragraph()
doc.add_heading('After (LlamaParse → MD):', 2)
doc.add_paragraph('## Adding Fractions with Different Denominators\n\nTo add fractions 1/2 + 1/3:\n1. Find the LCD of 2 and 3 = 6\n2. Convert: 1/2 = 3/6 and 1/3 = 2/6\n3. Add: 3/6 + 2/6 = 5/6')
doc.add_page_break()

# Slide 11
doc.add_heading('Slide 11: Smart Chunking with Metadata', 1)
doc.add_heading('From Dumb Chunks to Intelligent Blocks', 2)
table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Aspect'
hdr[1].text = 'Traditional'
hdr[2].text = 'Our Approach'
data = [
    ('Chunk Size', 'Fixed 512 tokens', 'Semantic boundaries'),
    ('Structure', 'None', 'Chapter → Section → Block'),
    ('Metadata', 'filename only', 'chapter, section, page, block_type'),
    ('Block Types', 'None', 'DEFINITION, PROCEDURE, EXAMPLE, PRACTICE'),
    ('Overlap', 'None', 'Context-aware overlap')
]
for i, (aspect, trad, our) in enumerate(data):
    table.rows[i+1].cells[0].text = aspect
    table.rows[i+1].cells[1].text = trad
    table.rows[i+1].cells[2].text = our
doc.add_paragraph()
doc.add_heading('Block Type Weighting System:', 2)
doc.add_paragraph('DEFINITION: 1.3 (Boost definitions)')
doc.add_paragraph('PROCEDURE: 1.2 (Boost how-to content)')
doc.add_paragraph('EXPLANATION: 1.1 (Boost explanations)')
doc.add_paragraph('BODY: 1.0 (Normal)')
doc.add_paragraph('WEBLINKS: 0.3 (Demote irrelevant)')
doc.add_page_break()

# Slide 12
doc.add_heading('Slide 12: Hybrid Retrieval - The Game Changer', 1)
doc.add_heading('Why Single Retrieval Method Fails', 2)
table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Query'
hdr[1].text = 'Dense Search'
hdr[2].text = 'BM25 Search'
hdr[3].text = 'Winner'
data = [
    ('"What is addition?"', '✅ Good', '❌ Misses', 'Dense'),
    ('"Exercise 2.3"', '❌ Fails', '✅ Exact match', 'BM25'),
    ('"How to solve sum?"', '⚠️ OK', '⚠️ OK', 'Both')
]
for i, (q, d, b, w) in enumerate(data):
    table.rows[i+1].cells[0].text = q
    table.rows[i+1].cells[1].text = d
    table.rows[i+1].cells[2].text = b
    table.rows[i+1].cells[3].text = w
doc.add_paragraph()
doc.add_heading('Our Hybrid Solution:', 2)
doc.add_paragraph('Math: 65% Dense + 35% BM25 + CrossEncoder Reranking')
doc.add_paragraph('Science/Computer: RRF Fusion + Keyword Gating')
doc.add_paragraph()
doc.add_heading('Reciprocal Rank Fusion (RRF):', 2)
items = [
    'Document in both lists → boosted score',
    'Document in one list → lower score',
    'Best of both worlds'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_page_break()

# Slide 13
doc.add_heading('Slide 13: Cross-Encoder Reranking', 1)
doc.add_heading('The Final Quality Filter', 2)
doc.add_paragraph('Problem: Initial retrieval returns ~40 candidates')
doc.add_paragraph('Need: Select the TOP 6 most relevant')
doc.add_paragraph()
doc.add_heading('Traditional Approach:', 2)
items = ['Use embedding similarity score', 'Fast but not accurate']
for item in items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()
doc.add_heading('Our Cross-Encoder Approach:', 2)
items = [
    'Neural model scores (query, document) pairs',
    'Much more accurate than embeddings',
    'Model: ms-marco-MiniLM-L-6-v2'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('Query: "How to multiply fractions?"')
doc.add_paragraph()
doc.add_paragraph('Before Reranking:')
doc.add_paragraph('1. "Fractions are parts..." (0.72) ← Wrong!')
doc.add_paragraph('2. "To multiply fractions, multiply numerators..." (0.68)')
doc.add_paragraph()
doc.add_paragraph('After Reranking:')
doc.add_paragraph('1. "To multiply fractions, multiply numerators..." (0.91) ✅')
doc.add_paragraph('2. "Fractions are parts..." (0.34)')
doc.add_page_break()

# Slide 14
doc.add_heading('Slide 14: Intent Detection - Understanding the Question', 1)
doc.add_heading('Different Questions Need Different Content', 2)
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Intent'
hdr[1].text = 'Example'
hdr[2].text = 'Best Block Types'
data = [
    ('DEFINITION', '"What is a fraction?"', 'GLOSSARY, DEFINITION'),
    ('PROBLEM', '"Solve 2/3 + 1/4"', 'PRACTICE, QUESTION'),
    ('EXAMPLE', '"Show me an example"', 'EXAMPLE, EXPLANATION'),
    ('PROCEDURE', '"How to add fractions?"', 'PROCEDURE, RULE')
]
for i, (intent, ex, blocks) in enumerate(data):
    table.rows[i+1].cells[0].text = intent
    table.rows[i+1].cells[1].text = ex
    table.rows[i+1].cells[2].text = blocks
doc.add_paragraph()
doc.add_heading('Our Intent Detection:', 2)
doc.add_paragraph('if "solve" or "calculate" in query: return "PROBLEM"')
doc.add_paragraph('if "example" or "show" in query: return "EXAMPLE"')
doc.add_paragraph('if "what is" or "define" in query: return "DEFINITION"')
doc.add_paragraph('else: return "EXPLAIN"')
doc.add_paragraph()
p = doc.add_paragraph('Result: Right content type for right question!')
p.runs[0].bold = True
doc.add_page_break()

# Slide 15
doc.add_heading('Slide 15: Science & Computer - Enhanced RAG', 1)
doc.add_heading('"What" Questions - Factual Retrieval Done Right', 2)
doc.add_paragraph('Traditional RAG:')
doc.add_paragraph('Query → Embed → Search → LLM')
doc.add_paragraph()
doc.add_paragraph('Enhanced RAG:')
doc.add_paragraph('Query → Intent → Hybrid Search → Dense + BM25 → RRF Fusion → Keyword Gate → Block Filter → Rerank → Top Results → LLM')
doc.add_paragraph()
doc.add_heading('Enhancements Applied:', 2)
items = [
    '✅ Hybrid retrieval (Dense + BM25)',
    '✅ Reciprocal Rank Fusion (RRF)',
    '✅ Keyword gating (filters irrelevant)',
    '✅ Block type weighting',
    '✅ Subject-specific intent detection',
    '✅ Relevance scoring & validation'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()
p = doc.add_paragraph('Improvement: 45% → 88% accuracy')
p.runs[0].bold = True
p.runs[0].font.color.rgb = RGBColor(34, 197, 94)
doc.add_page_break()

# Slide 16
doc.add_heading('Slide 16: Mathematics - RAG to RAT (Revolutionary)', 1)
doc.add_heading('The "How" Problem', 2)
doc.add_paragraph('Traditional RAG answers "What" questions:')
doc.add_paragraph('"What is addition?" → ✅ "Addition is combining numbers"')
doc.add_paragraph()
doc.add_paragraph('But fails at "How" questions:')
doc.add_paragraph('"How do I add 24 + 38?" → ❌ "Addition is combining..." (useless!)')
doc.add_paragraph()
doc.add_heading('RAT: Retrieval-Augmented Thinking', 2)
doc.add_paragraph('Student: "How to add 24 + 38?"')
doc.add_paragraph()
doc.add_paragraph('Step 1: Generate Chain-of-Thought (before retrieval)')
doc.add_paragraph('→ "Need to add ones: 4+8, then tens: 2+3, carry if needed"')
doc.add_paragraph()
doc.add_paragraph('Step 2: Use EACH thought as retrieval query')
doc.add_paragraph('→ Query 1: "adding ones place"')
doc.add_paragraph('→ Query 2: "carrying in addition"')
doc.add_paragraph('→ Query 3: "adding tens place"')
doc.add_paragraph()
doc.add_paragraph('Step 3: Verify each step against textbook')
doc.add_paragraph('Step 4: Generate verified, step-by-step answer')
doc.add_paragraph()
p = doc.add_paragraph('RAT = Think first, then retrieve, then verify')
p.runs[0].bold = True
doc.add_page_break()

# Slide 17
doc.add_heading('Slide 17: Auto Image Generation - RAT in Action', 1)
doc.add_heading('Visual Learning for Mathematics', 2)
doc.add_paragraph('The Problem: Text alone doesn\'t explain "How"')
doc.add_paragraph('The Solution: Auto-generate visual explanations')
doc.add_paragraph()
doc.add_paragraph('Student: "What is 2+4?"')
doc.add_paragraph('AI: "2+4 equals 6!"')
doc.add_paragraph()
doc.add_paragraph('Student: "How?"')
doc.add_paragraph('AI: [Generates image showing 2 apples + 4 apples = 6 apples]')
doc.add_paragraph('"Look! 2 apples plus 4 apples gives us 6 apples total!"')
doc.add_paragraph()
doc.add_heading('How It Works:', 2)
steps = [
    'Detect "how" pattern in question',
    'Get context from chat memory',
    'Generate structured image prompt',
    'Create educational visual (DALL-E)',
    'Display with text explanation'
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')
doc.add_paragraph()
p = doc.add_paragraph('This is RAT: Retrieve knowledge + Generate reasoning + Visual proof')
p.runs[0].bold = True
doc.add_page_break()

# Slide 18
doc.add_heading('Slide 18: Technologies Used - The Stack', 1)
doc.add_heading('Complete Tech Stack', 2)
table = doc.add_table(rows=11, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Layer'
hdr[1].text = 'Technology'
hdr[2].text = 'Purpose'
data = [
    ('LLM', 'GPT-4o-mini', 'Response generation'),
    ('Embeddings', 'SentenceTransformer (all-MiniLM-L6-v2)', 'Semantic search'),
    ('Vector DB', 'ChromaDB', 'Document storage'),
    ('Reranker', 'CrossEncoder (ms-marco-MiniLM-L-6-v2)', 'Quality filtering'),
    ('Sparse Search', 'BM25 (rank_bm25)', 'Keyword matching'),
    ('PDF Parsing', 'LlamaParse (LlamaIndex)', 'MD conversion'),
    ('Image Gen', 'DALL-E / GPT-Image', 'Visual aids'),
    ('TTS', 'OpenAI TTS', 'Audio explanations'),
    ('Frontend', 'Streamlit', 'Web interface'),
    ('Deployment', 'Streamlit Cloud', 'Free hosting')
]
for i, (layer, tech, purpose) in enumerate(data):
    table.rows[i+1].cells[0].text = layer
    table.rows[i+1].cells[1].text = tech
    table.rows[i+1].cells[2].text = purpose
doc.add_page_break()

# Slide 19
doc.add_heading('Slide 19: Why These Specific Technologies?', 1)
doc.add_heading('Every Choice Has a Reason', 2)
table = doc.add_table(rows=8, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Decision'
hdr[1].text = 'Why This'
hdr[2].text = 'Not That'
data = [
    ('GPT-4o-mini', '10x cheaper, fast, smart enough', 'GPT-4 (expensive, slow)'),
    ('SentenceTransformer', 'Free, local, fast', 'OpenAI Embeddings (paid)'),
    ('ChromaDB', 'Free, easy, persistent', 'Pinecone (paid cloud)'),
    ('BM25 + Dense', 'Hybrid catches more', 'Dense only (misses keywords)'),
    ('CrossEncoder', 'High accuracy reranking', 'Embedding similarity (weak)'),
    ('LlamaParse', 'Preserves structure', 'OCR (destroys math)'),
    ('Streamlit', 'Rapid dev, clean UI', 'React (complex, slow dev)')
]
for i, (dec, why, not_that) in enumerate(data):
    table.rows[i+1].cells[0].text = dec
    table.rows[i+1].cells[1].text = why
    table.rows[i+1].cells[2].text = not_that
doc.add_paragraph()
p = doc.add_paragraph('Philosophy: Best tool for the job, not the most expensive')
p.runs[0].bold = True
doc.add_page_break()

# Slide 20
doc.add_heading('Slide 20: Key Features Implemented', 1)
doc.add_heading('Complete Feature Set', 2)
table = doc.add_table(rows=11, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Feature'
hdr[1].text = 'What It Does'
hdr[2].text = 'Technology'
data = [
    ('AI Tutor Chat', 'Patient, step-by-step explanations', 'GPT-4o-mini + RAG'),
    ('Auto Image Gen', 'Visual aids for "how" questions', 'DALL-E + RAT'),
    ('Text-to-Speech', 'Audio playback of answers', 'OpenAI TTS'),
    ('Bilingual UI', 'Full English & Urdu support', 'i18n system'),
    ('Practice Quizzes', 'AI-generated questions', 'GPT-4o-mini'),
    ('Learning Analytics', 'Progress charts & stats', 'Plotly'),
    ('Chat Memory', 'Context-aware conversations', 'Session state'),
    ('Subject Validation', 'Blocks off-topic questions', 'Keyword + relevance'),
    ('Curriculum Aligned', 'Answers from textbooks', 'RAG pipeline'),
    ('Rewards System', 'Stars & encouragement', 'Gamification')
]
for i, (feat, what, tech) in enumerate(data):
    table.rows[i+1].cells[0].text = feat
    table.rows[i+1].cells[1].text = what
    table.rows[i+1].cells[2].text = tech
doc.add_page_break()

# Slide 21
doc.add_heading('Slide 21: Evaluation Results - The Numbers Speak', 1)
doc.add_heading('Head-to-Head Comparison', 2)
table = doc.add_table(rows=7, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Metric'
hdr[1].text = 'OCR + Basic RAG'
hdr[2].text = 'LlamaParse + Enhanced RAG'
hdr[3].text = 'Improvement'
data = [
    ('Text Extraction Quality', '45%', '98%', '+53%'),
    ('Math Symbol Accuracy', '30%', '95%', '+65%'),
    ('Retrieval Precision@5', '42%', '87%', '+45%'),
    ('Answer Relevance', '55%', '91%', '+36%'),
    ('Out-of-Textbook Detection', '20%', '85%', '+65%'),
    ('Response Latency', '4.2s', '1.8s', '-57%')
]
for i, (metric, ocr, enhanced, imp) in enumerate(data):
    table.rows[i+1].cells[0].text = metric
    table.rows[i+1].cells[1].text = ocr
    table.rows[i+1].cells[2].text = enhanced
    table.rows[i+1].cells[3].text = imp
doc.add_page_break()

# Slide 22
doc.add_heading('Slide 22: Real Query Comparison', 1)
doc.add_heading('Same Questions, Different Results', 2)
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Student Query'
hdr[1].text = 'OCR + Basic RAG'
hdr[2].text = 'Enhanced RAG'
data = [
    ('"How to add 1/2 + 1/3?"', '"Fractions are parts of a whole..." ❌', '"Step 1: Find LCD=6, Step 2: Convert..." ✅'),
    ('"What is √16?"', '"V16 = ?" (OCR error) ❌', '"√16 = 4 because 4×4=16" ✅'),
    ('"Solve Exercise 2.4"', 'Random chapter content ❌', 'Exact exercise with solution ✅'),
    ('"What is the Internet?"', 'Generic definition ❌', 'Textbook definition + examples ✅')
]
for i, (query, ocr, enhanced) in enumerate(data):
    table.rows[i+1].cells[0].text = query
    table.rows[i+1].cells[1].text = ocr
    table.rows[i+1].cells[2].text = enhanced
doc.add_page_break()

# Slide 23
doc.add_heading('Slide 23: Why Our System is Better', 1)
doc.add_heading('The Complete Picture', 2)
doc.add_paragraph('BEFORE (Traditional RAG):')
items = [
    '❌ OCR destroys math',
    '❌ Fixed chunk size',
    '❌ Single retrieval (dense only)',
    '❌ No reranking',
    '❌ No intent detection',
    '❌ No subject validation',
    '❌ Same pipeline for all',
    '❌ Can\'t detect off-topic',
    '❌ No visual for "how"',
    '❌ Text only'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('AFTER (Enhanced RAG):')
items = [
    '✅ LlamaParse preserves all',
    '✅ Semantic chunking + metadata',
    '✅ Hybrid (Dense + BM25 + RRF)',
    '✅ Cross-Encoder reranking',
    '✅ Intent-aware block filtering',
    '✅ Subject keyword checking',
    '✅ Subject-specific pipelines',
    '✅ Relevance scoring + thresholds',
    '✅ Auto image generation for math',
    '✅ Multi-modal (text + image + voice)'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()
p = doc.add_paragraph('Bottom Line: We didn\'t just improve RAG, we reinvented it for education')
p.runs[0].bold = True
doc.add_page_break()

# Slide 24
doc.add_heading('Slide 24: Autism-Friendly Design', 1)
doc.add_heading('Every Feature Designed for Special Needs', 2)
table = doc.add_table(rows=9, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Need'
hdr[1].text = 'Our Solution'
data = [
    ('Overwhelm easily', 'Clean, calm interface with soft blues'),
    ('Need clarity', 'Simple language, no jargon'),
    ('Visual learners', 'Auto-generated images for concepts'),
    ('Audio preference', 'Text-to-Speech for all answers'),
    ('Need patience', 'AI never rushes, always encouraging'),
    ('Need consistency', 'Same layout across all pages'),
    ('Need structure', 'Step-by-step explanations always'),
    ('Need encouragement', 'Stars, rewards, positive feedback')
]
for i, (need, solution) in enumerate(data):
    table.rows[i+1].cells[0].text = need
    table.rows[i+1].cells[1].text = solution
doc.add_paragraph()
p = doc.add_paragraph('Design Philosophy: "Reduce anxiety, increase understanding"')
p.runs[0].bold = True
doc.add_page_break()

# Slide 25
doc.add_heading('Slide 25: Future Work - Agentic AI', 1)
doc.add_heading('The Next Evolution: Autonomous Agents', 2)
doc.add_paragraph('Current System:')
items = ['Rules-based response type selection', 'Manual triggers for image/voice']
for item in items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('Agentic AI System:')
items = ['Autonomous decision making', 'Learns student preferences', 'Proactive suggestions']
for item in items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()
doc.add_heading('Planned Agent Architecture:', 2)
doc.add_paragraph('Orchestrator Agent (Coordinates all decisions)')
doc.add_paragraph('├── Content Agent: What to retrieve')
doc.add_paragraph('├── Media Agent: Text vs Image vs Voice')
doc.add_paragraph('└── Adaptation Agent: Adjust difficulty')
doc.add_page_break()

# Slide 26
doc.add_heading('Slide 26: Agentic AI - The Vision', 1)
doc.add_heading('Intelligent Response Selection', 2)
doc.add_paragraph('Student: "How does multiplication work?"')
doc.add_paragraph()
doc.add_paragraph('Orchestrator Agent analyzes...')
doc.add_paragraph()
doc.add_paragraph('Content Agent:')
doc.add_paragraph('→ Detects: PROCEDURE intent')
doc.add_paragraph('→ Retrieves: multiplication steps from textbook')
doc.add_paragraph()
doc.add_paragraph('Media Agent decides:')
doc.add_paragraph('→ Text: YES (explain concept)')
doc.add_paragraph('→ Image: YES (visual demonstration)')
doc.add_paragraph('→ Voice: OFFER (student prefers reading)')
doc.add_paragraph()
doc.add_paragraph('Adaptation Agent:')
doc.add_paragraph('→ Notes: Student is visual learner')
doc.add_paragraph('→ Future: Auto-prioritize images')
doc.add_paragraph()
doc.add_paragraph('Final Response:')
doc.add_paragraph('✅ Text explanation with simple steps')
doc.add_paragraph('✅ Auto-generated multiplication visual')
doc.add_paragraph('✅ "Would you like me to read this?" button')
doc.add_page_break()

# Slide 27
doc.add_heading('Slide 27: Project Roadmap', 1)
doc.add_heading('What We\'ve Achieved & What\'s Next', 2)
table = doc.add_table(rows=9, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Phase'
hdr[1].text = 'Status'
hdr[2].text = 'Features'
data = [
    ('Phase 1', '✅ Complete', 'OCR-based RAG (failed)'),
    ('Phase 2', '✅ Complete', 'LlamaParse + Enhanced RAG'),
    ('Phase 3', '✅ Complete', 'RAT for Mathematics'),
    ('Phase 4', '✅ Complete', 'Auto Image + TTS + Bilingual'),
    ('Phase 5', '✅ Complete', 'Practice Quizzes + Analytics'),
    ('Phase 6', '🔄 In Progress', 'User testing with students'),
    ('Phase 7', '📋 Planned', 'Agentic AI implementation'),
    ('Phase 8', '📋 Planned', 'More subjects, Mobile app')
]
for i, (phase, status, features) in enumerate(data):
    table.rows[i+1].cells[0].text = phase
    table.rows[i+1].cells[1].text = status
    table.rows[i+1].cells[2].text = features
doc.add_page_break()

# Slide 28
doc.add_heading('Slide 28: Summary - What We Built', 1)
doc.add_heading('The Complete Picture', 2)
table = doc.add_table(rows=8, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Challenge'
hdr[1].text = 'Our Solution'
hdr[2].text = 'Result'
data = [
    ('OCR destroys math', 'LlamaParse to MD', '98% accuracy'),
    ('Weak retrieval', 'Hybrid + Reranking', '87% precision'),
    ('No "how" support', 'RAT + Auto Images', 'Visual learning'),
    ('English only', 'Bilingual system', 'Urdu + English'),
    ('No assessment', 'Practice quizzes', 'AI-generated tests'),
    ('No tracking', 'Learning analytics', 'Charts & progress'),
    ('Not autism-friendly', 'Calm, patient design', 'Reduced anxiety')
]
for i, (challenge, solution, result) in enumerate(data):
    table.rows[i+1].cells[0].text = challenge
    table.rows[i+1].cells[1].text = solution
    table.rows[i+1].cells[2].text = result
doc.add_paragraph()
p = doc.add_paragraph('We transformed a failing RAG into an intelligent tutoring system')
p.runs[0].bold = True
doc.add_page_break()

# Slide 29
doc.add_heading('Slide 29: Conclusion', 1)
doc.add_heading('Key Takeaways', 2)
doc.add_paragraph()
doc.add_paragraph('1. Problem: Traditional RAG fails for educational content')
items = ['OCR destroys structure', 'No intent understanding', 'Can\'t handle "how" questions']
for item in items:
    doc.add_paragraph('   • ' + item)
doc.add_paragraph()
doc.add_paragraph('2. Innovation: Subject-specific enhanced pipelines')
items = ['LlamaParse preserves content', 'Hybrid retrieval catches more', 'RAT for procedural math']
for item in items:
    doc.add_paragraph('   • ' + item)
doc.add_paragraph()
doc.add_paragraph('3. Impact: Education for students who need it most')
items = ['Autism-friendly design', 'Bilingual support', 'Visual + Audio learning']
for item in items:
    doc.add_paragraph('   • ' + item)
doc.add_paragraph()
p = doc.add_paragraph('"Every child deserves education tailored to their unique needs"')
p.runs[0].bold = True
p.runs[0].italic = True
doc.add_page_break()

# Slide 30
doc.add_heading('Slide 30: Thank You', 1)
doc.add_paragraph()
title = doc.add_paragraph('🎓 AutiStudy')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].bold = True
title.runs[0].font.size = Pt(28)
title.runs[0].font.color.rgb = RGBColor(37, 99, 235)

subtitle = doc.add_paragraph('AI-Powered Adaptive Learning')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('for Students with Autism')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p = doc.add_paragraph('Live Demo: Streamlit Cloud')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('GitHub: github.com/zk-007/AutiStudy')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph('Built with ❤️ for students in Pakistan')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True

doc.add_paragraph()
p = doc.add_paragraph('Thank you for your attention!')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('Questions?')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True
p.runs[0].font.size = Pt(24)

# Save
doc.save('AutiStudy_Presentation.docx')
print("Document saved as AutiStudy_Presentation.docx")
