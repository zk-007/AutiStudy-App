"""
Build the Agent Evaluation Word document from the eval result JSON files.
=========================================================================

Reads:
  scripts/eval/results/router_eval.json   (deterministic visual-aid router)
  scripts/eval/results/agent_eval.json     (GPT-4o ReAct teaching agent)

Writes:
  docs/AutiStudy_Agent_Evaluation.docx

Run (after running the two eval scripts) from AutiStudy-React:
  python scripts/eval/router_eval.py
  python scripts/eval/agent_eval.py --k 5
  python scripts/generate_agent_evaluation_docx.py

Requires: python-docx
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

REACT_ROOT = Path(__file__).resolve().parent.parent
RESULTS_DIR = REACT_ROOT / "scripts" / "eval" / "results"
ROUTER_JSON = RESULTS_DIR / "router_eval.json"
AGENT_JSON = RESULTS_DIR / "agent_eval.json"
OUT_PATH = REACT_ROOT / "docs" / "AutiStudy_Agent_Evaluation.docx"

NAVY = RGBColor(15, 45, 74)


def style_doc(doc: Document) -> None:
    for name, size in (("Heading 1", 16), ("Heading 2", 13)):
        st = doc.styles[name]
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = NAVY


def add_metric_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for ri, (k, v) in enumerate(rows):
        table.rows[ri].cells[0].text = k
        table.rows[ri].cells[1].text = v
    doc.add_paragraph()


def add_per_class_table(doc: Document, per_class: dict) -> None:
    labels = sorted(per_class)
    table = doc.add_table(rows=1 + len(labels), cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(["Class / route", "Precision", "Recall", "F1", "Support"]):
        table.rows[0].cells[i].text = h
    for ri, lbl in enumerate(labels):
        s = per_class[lbl]
        cells = table.rows[ri + 1].cells
        cells[0].text = lbl
        cells[1].text = f"{s['precision']:.2f}"
        cells[2].text = f"{s['recall']:.2f}"
        cells[3].text = f"{s['f1']:.2f}"
        cells[4].text = str(int(s["support"]))
    doc.add_paragraph()


def build() -> Document:
    router = json.loads(ROUTER_JSON.read_text(encoding="utf-8"))
    agent = json.loads(AGENT_JSON.read_text(encoding="utf-8"))

    doc = Document()
    style_doc(doc)

    doc.add_heading("AutiStudy — Agent Evaluation", 0)
    doc.add_paragraph(
        "This report evaluates AutiStudy's adaptive teaching agent on the two "
        "metrics requested for agent evaluation: Routing Accuracy (does the "
        "agent send each input to the correct action/branch?) and Response "
        "Consistency (does the agent make the same decision for the same input?). "
        "Two routing layers are assessed: (1) the deterministic visual-aid "
        "router and (2) the GPT-4o ReAct teaching agent."
    )

    # ── 1. Methodology ──────────────────────────────────────────────────────
    doc.add_heading("1. Methodology", 1)
    doc.add_paragraph(
        "Routing Accuracy is measured against gold-labeled inputs, where each "
        "label is the pedagogically-correct route. We report overall accuracy, "
        "per-class precision / recall / F1, and macro / weighted means so that "
        "rare routes are not hidden by common ones.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Response Consistency is measured by repeating each input k times and "
        "quantifying decision stability: mean majority agreement, all-agree "
        "rate, Fleiss' kappa, and normalized entropy (0 = perfectly stable).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "The deterministic router needs no model calls (regex), so its response "
        "consistency is exact. The ReAct agent is exercised with its real tool "
        "schema, system prompt, model (gpt-4o) and temperature (0.2), with "
        "cross-session memory held at the neutral first-session default so each "
        "repeat sees an identical input.",
        style="List Bullet",
    )

    # ── 2. Visual-aid router ────────────────────────────────────────────────
    doc.add_heading("2. Visual-Aid Router (deterministic)", 1)
    doc.add_paragraph(
        f"System under test: {router['system_under_test']} "
        f"({router['router_type']}). The router classifies a student question "
        f"into one of 11 visual tracks."
    )
    doc.add_heading("2.1 Routing accuracy", 2)
    add_metric_table(doc, [
        ("Items evaluated", str(router["n_items"])),
        ("Routing accuracy", f"{router['routing_accuracy']:.1%}"),
        ("Macro F1", f"{router['macro']['f1']:.3f}"),
        ("Weighted F1", f"{router['weighted']['f1']:.3f}"),
        ("Macro precision", f"{router['macro']['precision']:.3f}"),
        ("Macro recall", f"{router['macro']['recall']:.3f}"),
    ])
    doc.add_heading("2.2 Per-route breakdown", 2)
    add_per_class_table(doc, router["per_class"])

    doc.add_heading("2.3 Response consistency", 2)
    rc = router["response_consistency"]
    doc.add_paragraph(
        f"{rc['note']} Verified identical on repeat: "
        f"{rc['verified_identical_on_repeat']} — agreement rate "
        f"{rc['agreement_rate']:.0%}."
    )

    if router.get("failures"):
        doc.add_heading("2.4 Misroutes found", 2)
        doc.add_paragraph(
            "Routing-accuracy evaluation surfaced the following genuine "
            "misroutes (useful for fixing the router):"
        )
        for f in router["failures"]:
            doc.add_paragraph(
                f"\"{f['question']}\" — expected {f['expected']}, got {f['predicted']}.",
                style="List Bullet",
            )

    # ── 3. ReAct teaching agent ─────────────────────────────────────────────
    doc.add_heading("3. ReAct Teaching Agent (GPT-4o)", 1)
    doc.add_paragraph(
        f"System under test: {agent['system_under_test']}. Model "
        f"{agent['model']} at temperature {agent['temperature']}. The agent "
        f"selects one teaching action from 8 tools based on the student's "
        f"emotion and confusion state."
    )
    ra = agent["routing_accuracy"]
    doc.add_heading("3.1 Routing accuracy", 2)
    add_metric_table(doc, [
        ("Scenarios", str(agent["n_scenarios"])),
        ("Repeats per scenario (k)", str(agent["k_repeats"])),
        ("Total model calls", str(agent["total_model_calls"])),
        ("Routing accuracy (strict, single best route)", f"{ra['strict_primary']:.1%}"),
        ("Routing accuracy (lenient, accepted set)", f"{ra['lenient_accepted_set']:.1%}"),
        ("Macro F1", f"{ra['macro']['f1']:.3f}"),
    ])
    doc.add_heading("3.2 Per-action breakdown", 2)
    add_per_class_table(doc, ra["per_class"])

    doc.add_heading("3.3 Response consistency", 2)
    cons = agent["response_consistency"]
    add_metric_table(doc, [
        ("Mean majority agreement", f"{cons['mean_majority_agreement']:.1%}"),
        (f"All-{cons['k_repeats']}-agree rate", f"{cons['all_agree_rate']:.1%}"),
        ("Fleiss' kappa", f"{cons['fleiss_kappa']:.3f}"),
        ("Mean entropy (0 = perfectly stable)", f"{cons['mean_entropy']:.3f}"),
    ])

    # ── 4. Summary ──────────────────────────────────────────────────────────
    doc.add_heading("4. Summary", 1)
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    hdr = ["Component", "Routing accuracy", "Response consistency"]
    for i, h in enumerate(hdr):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].text = "Visual-aid router (deterministic)"
    table.rows[1].cells[1].text = f"{router['routing_accuracy']:.1%}"
    table.rows[1].cells[2].text = "100% (deterministic)"
    table.rows[2].cells[0].text = "ReAct teaching agent (GPT-4o)"
    table.rows[2].cells[1].text = (
        f"{ra['strict_primary']:.1%} strict / {ra['lenient_accepted_set']:.1%} lenient"
    )
    table.rows[2].cells[2].text = (
        f"{cons['mean_majority_agreement']:.0%} agreement, "
        f"kappa {cons['fleiss_kappa']:.2f}"
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Recommended headline metrics: (a) overall Routing Accuracy with macro "
        "F1, and (b) mean majority agreement with Fleiss' kappa for Response "
        "Consistency. These mirror the retriever/RAT accuracy already reported "
        "and give a complete, defensible agent-evaluation story."
    )
    return doc


def main() -> None:
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build()
    doc.save(str(OUT_PATH))
    print(f"Created: {OUT_PATH}")


if __name__ == "__main__":
    main()
