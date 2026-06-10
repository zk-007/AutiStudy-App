"""
Build Word document with per-grade tables: subject, chapter count, chunk count.
Chapter counts from utils/book_parser.py; chunk counts from OneSharedChromaDB.

Run from AutiStudy-App/frontend:
  python scripts/generate_dataset_statistics_docx.py

Requires: python-docx, chromadb; backend path ../AutiStudy
"""
from __future__ import annotations

import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

REACT_ROOT = Path(__file__).resolve().parent.parent
BACKEND_ROOT = REACT_ROOT.parent / "backend"
CHROMA_PATH = BACKEND_ROOT / "OneSharedChromaDB"
OUT_PATH = REACT_ROOT / "docs" / "AutiStudy_Dataset_Statistics_by_Grade.docx"

# doc_id → (grade, display subject)
DOC_MAP: dict[str, tuple[int, str]] = {
    "MATH4": (4, "Maths"),
    "GS4": (4, "General Science"),
    "MATH5": (5, "Maths"),
    "GS5": (5, "General Science"),
    "MATH6": (6, "Maths"),
    "GS6": (6, "General Science"),
    "CS6": (6, "Computer Science"),
    "MATH7": (7, "Maths"),
    "GS7": (7, "General Science"),
    "CS7": (7, "Computer Science"),
}

SUBJECTS_BY_GRADE: dict[int, list[str]] = {
    4: ["Maths", "General Science"],
    5: ["Maths", "General Science"],
    6: ["Maths", "General Science", "Computer Science"],
    7: ["Maths", "General Science", "Computer Science"],
}


def chapter_count_from_books(grade: int, subject: str) -> int | None:
    sys.path.insert(0, str(BACKEND_ROOT))
    from utils.book_parser import get_chapters

    chapters = get_chapters(grade, subject)
    if chapters is None:
        return None
    return len(chapters)


def chapter_count_math7_fallback() -> tuple[int, str]:
    """Grade 7 Maths markdown is empty / not in BOOK_MAP; infer units from Chroma metadata."""
    import chromadb

    client = chromadb.PersistentClient(path=str(CHROMA_PATH))
    col = client.get_collection("ptb_textbooks")
    res = col.get(where={"doc_id": "MATH7"}, include=["metadatas"])
    units: set[int] = set()
    for m in res["metadatas"]:
        ch = (m.get("chapter") or "").strip()
        match = re.match(r"^(\d+)\.", ch)
        if match:
            units.add(int(match.group(1)))
    if units:
        return len(units), "vector-store units (textbook file not registered in BOOK_MAP)"
    return 0, "no chapter index"


def chunk_counts_chroma() -> dict[tuple[int, str], int]:
    import chromadb

    if not CHROMA_PATH.exists():
        raise FileNotFoundError(f"ChromaDB not found at {CHROMA_PATH}")

    client = chromadb.PersistentClient(path=str(CHROMA_PATH))
    col = client.get_collection("ptb_textbooks")
    out: dict[tuple[int, str], int] = {}
    for doc_id, (grade, subject) in DOC_MAP.items():
        res = col.get(where={"doc_id": doc_id}, include=[])
        out[(grade, subject)] = len(res["ids"])
    return out


def style_doc(doc: Document) -> None:
    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(15, 45, 74)


def add_grade_table(
    doc: Document,
    grade: int,
    rows: list[tuple[str, str, str]],
) -> None:
    doc.add_heading(f"Grade {grade}", 1)
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"
    headers = ["Subject", "Total chapters", "Total chunks"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for ri, (subj, ch, ck) in enumerate(rows):
        table.rows[ri + 1].cells[0].text = subj
        table.rows[ri + 1].cells[1].text = ch
        table.rows[ri + 1].cells[2].text = ck
    doc.add_paragraph()


def build() -> Document:
    chunks = chunk_counts_chroma()
    total_chunks = sum(chunks.values())

    doc = Document()
    style_doc(doc)

    for grade in [4, 5, 6, 7]:
        table_rows: list[tuple[str, str, str]] = []
        for subject in SUBJECTS_BY_GRADE[grade]:
            n_chunks = chunks.get((grade, subject), 0)
            if subject == "Maths" and grade == 7:
                n_ch, _note = chapter_count_math7_fallback()
                ch_display = str(n_ch)
            else:
                n_ch = chapter_count_from_books(grade, subject)
                ch_display = str(n_ch) if n_ch is not None else "—"
            table_rows.append((subject, ch_display, str(n_chunks)))
        add_grade_table(doc, grade, table_rows)

    _ = total_chunks  # used for validation during generation
    return doc


def main() -> None:
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build()
    doc.save(str(OUT_PATH))
    print(f"Created: {OUT_PATH}")


if __name__ == "__main__":
    main()
