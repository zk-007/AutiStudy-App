"""
Generate Experimental Setup Word document (v16 — concise, cmt3-compliant).

cmt3 rule: minimal top info; A = 3-4 sentences, D = 4-5 sentences, F = 3-4 sentences.

Run from frontend/:
  python scripts/generate_experimental_setup_docx.py

Output:
  docs/AutiStudy_Experimental_Setup_v16.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

REACT_ROOT = Path(__file__).resolve().parent.parent
OUT_PATH = REACT_ROOT / "docs" / "AutiStudy_Experimental_Setup_v16.docx"


def build() -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.size = Pt(11)
    normal.font.name = "Times New Roman"

    # ── A. Development and Evaluation Environment (3 sentences) ─────────
    doc.add_heading("A. Development and Evaluation Environment", level=2)
    doc.add_paragraph(
        "Experiments ran on an Intel Core i5-8365U CPU with 16 GB RAM under "
        "Windows 11 Pro (x64). The backend used Python 3.11.9 with FastAPI, "
        "ChromaDB, sentence-transformers (all-MiniLM-L6-v2), rank-bm25, and "
        "OpenAI SDK (gpt-4o-mini / gpt-4o); CV evaluation used Google Chrome "
        "with in-browser MediaPipe and face-api.js. No GPU training or "
        "fine-tuning was performed—all models were used as pre-trained components."
    )

    # ── D. Evaluation Query Set Construction (4 sentences) ────────────────
    doc.add_heading("D. Evaluation Query Set Construction", level=2)
    doc.add_paragraph(
        "Thirty-five core queries were built: 20 procedural maths (five per "
        "grade, 4–7), five maths retrieval items (Grade 4), and ten Grade 6 "
        "Computer Science items, plus comparable General Science sets for "
        "Grades 4, 6, and 7. Queries covered four types—factual, conceptual, "
        "procedural, and multi-step—and spanned Maths (4–7), General Science "
        "(4–7), and Computer Science (6–7). For maths, two to four keyword "
        "phrases per query defined gold relevance (automated match in retrieved "
        "text). For Computer Science and General Science, the project team "
        "manually labelled top-10 retrieved chunks against the textbook to "
        "compute Precision@5, Recall@5, MRR, and nDCG@5."
    )

    # ── F. Agent Evaluation Setup (3 sentences) ─────────────────────────────
    doc.add_heading("F. Agent Evaluation Setup", level=2)
    doc.add_paragraph(
        "The visual-aid router was evaluated on 60 gold-labelled inputs covering "
        "all eleven visual aid types; gold labels followed pedagogical guidelines "
        "from teacher consultation. The Media Agent was tested on 12 scenarios "
        "with five repeats each at temperature 0.2, memory held at the neutral "
        "first-session default so every repeat received identical input context. "
        "Routing accuracy and response consistency were computed from the saved "
        "evaluation JSON outputs."
    )

    return doc


def main() -> None:
    doc = build()
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
