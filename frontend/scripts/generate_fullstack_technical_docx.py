"""
Generate a single AutiStudy full-stack technical DOCX (frontend + backend).
Run from repo root or frontend/:
  python frontend/scripts/generate_fullstack_technical_docx.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT_PATH = Path(__file__).resolve().parents[2] / "AutiStudy_FullStack_Technical_Documentation.docx"


def style_doc(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(2, 132, 199)
    h2 = doc.styles["Heading 2"]
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(15, 45, 74)
    h3 = doc.styles["Heading 3"]
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(30, 64, 100)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = cell_text
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.style = "Intense Quote"
    for run in p.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def build() -> Document:
    doc = Document()
    style_doc(doc)

    title = doc.add_heading("AutiStudy — Full-Stack Technical Documentation", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        "Frontend (Next.js) + Backend (FastAPI) · Production architecture · v6"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph(
        "Deploy targets: frontend/ → Vercel · backend/ → Railway · "
        "Release snapshot: releases/comprehension-production-stable-v6/"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── 1 Overview ──────────────────────────────────────────────────────────
    doc.add_heading("1. System Overview", 1)
    doc.add_paragraph(
        "AutiStudy is an autism-friendly adaptive tutoring platform for Pakistani "
        "curriculum (grades 4–7). The browser client is a Next.js 14 App Router app; "
        "the API is a FastAPI service that owns authentication, RAG-grounded tutoring, "
        "quiz generation, parent/family linking, email OTP, and an emotion-aware "
        "adaptive teaching agent."
    )
    code(
        doc,
        "Browser (Next.js 14)\n"
        "   │  HTTPS JSON + Bearer token\n"
        "   ▼\n"
        "FastAPI (api_server.py)\n"
        "   ├── data/*.json          (users, sessions, chats, OTPs, family)\n"
        "   ├── quiz_data/           (per-user quiz attempts)\n"
        "   ├── OneSharedChromaDB/   (textbook embeddings)\n"
        "   ├── OpenAI               (chat, quiz, TTS, images, agent)\n"
        "   └── SMTP                 (OTP, password reset, contact, family mail)",
    )

    doc.add_heading("1.1 Production Topology", 2)
    table(
        doc,
        ["Layer", "Platform", "Source folder", "Public URL (typical)"],
        [
            [
                "Frontend",
                "Vercel",
                "frontend/",
                "https://auti-study-app.vercel.app",
            ],
            [
                "Backend",
                "Railway",
                "backend/ (Dockerfile)",
                "https://autistudy-app-production.up.railway.app",
            ],
        ],
    )
    bullets(
        doc,
        [
            "Local: frontend npm run dev:all → Next :3000 + uvicorn :8000",
            "CORS: localhost:3000/3001 + CORS_ORIGINS + *.vercel.app regex",
            "Health: GET /api/health",
        ],
    )

    # ── 2 Frontend ──────────────────────────────────────────────────────────
    doc.add_heading("2. Frontend Technical Architecture", 1)

    doc.add_heading("2.1 Stack", 2)
    table(
        doc,
        ["Technology", "Role"],
        [
            ["Next.js 14.2 (App Router)", "SSR/CSR pages, routing, build"],
            ["React 18 + TypeScript 5", "UI components and typed client"],
            ["Tailwind CSS 3.4", "Design tokens (glacier/mint/deep), darkMode: class"],
            ["framer-motion / gsap / lenis", "Motion and smooth scroll"],
            ["KaTeX + react-markdown", "Math and rich tutor answers"],
            ["face-api.js / MediaPipe", "Optional expression / camera labs"],
            ["lucide-react", "Icons"],
        ],
    )

    doc.add_heading("2.2 Repository Layout (frontend)", 2)
    code(
        doc,
        "frontend/\n"
        "├── app/                    # App Router pages\n"
        "│   ├── layout.tsx          # Providers + theme boot script\n"
        "│   ├── globals.css         # CSS variables, dark theme, .bg-cta\n"
        "│   ├── dashboard/          # Child home\n"
        "│   ├── chat/               # Tutor + adaptive agent UI\n"
        "│   ├── quiz/, analytics/\n"
        "│   ├── login/, signup/, forgot-password/\n"
        "│   ├── parent/dashboard/, parent/report/\n"
        "│   ├── family/respond/     # Email invite approve/reject\n"
        "│   ├── manual/, contact/, faq/, about/\n"
        "│   └── onboarding/learning-profile/\n"
        "├── components/             # UI, agent, settings, dashboard\n"
        "├── lib/\n"
        "│   ├── api/client.ts       # Typed REST client\n"
        "│   ├── auth/               # AuthProvider, parent session\n"
        "│   ├── settings/           # Theme, TTS, a11y\n"
        "│   └── i18n/               # en | ur + RTL\n"
        "└── books_mds/              # Curriculum MD (also used by backend Docker)",
    )

    doc.add_heading("2.3 Routes", 2)
    table(
        doc,
        ["Route", "Purpose"],
        [
            ["/", "Marketing landing"],
            ["/signup, /login, /forgot-password", "Auth (child | parent via ?role=)"],
            ["/onboarding/learning-profile", "One-time learning preferences"],
            ["/dashboard", "Child home: mood, streak, subjects, Study CTAs"],
            ["/chat", "RAG tutor + modality ladder + optional camera consent"],
            ["/quiz", "MCQ quizzes from chapters or chat"],
            ["/analytics", "Accuracy, time, subject breakdown"],
            ["/parent/dashboard, /parent/report", "Linked-child progress views"],
            ["/family/respond", "Tokenized email approve/reject for invites"],
            ["/manual, /contact, /faq, /about", "Help and support surfaces"],
            ["/settings (modal)", "Appearance, voice, language, account, family"],
        ],
    )

    doc.add_heading("2.4 Auth & Session (client)", 2)
    bullets(
        doc,
        [
            "Child bearer: localStorage keys autistudy_token + autistudy_session",
            "Parent bearer: autistudy_parent_token (separate from child)",
            "AuthProvider restores session via GET /api/auth/me",
            "Signup: POST child/parent signup → email OTP → POST /api/auth/verify-email",
            "Forgot password: request → OTP → verify → reset (authForgotApi)",
            "API base: NEXT_PUBLIC_API_URL, else http://127.0.0.1:8000 locally, "
            "else Railway production fallback in client.ts",
        ],
    )

    doc.add_heading("2.5 Settings, Theme, i18n", 2)
    bullets(
        doc,
        [
            "SettingsContext: theme light|dark, fontSize, reduceMotion, highContrast, "
            "focusMode, TTS (autoRead, speed, volume, narrator)",
            "Theme persisted device-wide in autistudy_theme (+ guest blob sync); "
            "layout.tsx boot script applies html.dark before paint",
            "LocaleProvider: en | ur, RTL for Urdu (autistudy_locale)",
            "Primary CTA utility classes .bg-cta / .bg-cta-r avoid to-deep gradients "
            "(--deep flips to light text in dark mode)",
        ],
    )

    doc.add_heading("2.6 API Client Surface", 2)
    doc.add_paragraph(
        "Primary module: frontend/lib/api/client.ts. Exports typed namespaces that "
        "map 1:1 to FastAPI route groups."
    )
    table(
        doc,
        ["Client export", "Backend prefix / notes"],
        [
            ["authApi / authVerifyApi", "/api/auth/* signup, login, OTP verify"],
            ["authForgotApi", "/api/auth/forgot-password/*"],
            ["userApi", "/api/users/me/* dashboard, mood, subjects, stats"],
            ["profileApi", "/api/profile/learning"],
            ["chatApi", "/api/chat/* sessions, messages, speech, images"],
            ["quizApi", "/api/quiz/*, /api/analytics"],
            ["parentApi", "/api/auth/parent/*, /api/parent/*"],
            ["familyApi", "/api/family/* invites and linking"],
            ["contactApi", "POST /api/contact"],
        ],
    )
    doc.add_paragraph(
        "Adaptive agent hooks (useComprehensionFlow / useAdaptiveTutorAgent) often "
        "call /api/agent/* directly with fetch(API_BASE + …)."
    )

    doc.add_heading("2.7 Key UX Features", 2)
    bullets(
        doc,
        [
            "Dashboard subject cards → Study CTA opens chat for that subject",
            "Four explanation modalities: text, step-by-step, picture, voice",
            "Thumbs feedback drives modality ladder; calming break after repeated fails",
            "Optional camera (consent) for engagement/comfort signals — no video stored",
            "Parent family invites (FAM-XXXXX), max two linked parents",
            "Contact form emails support inbox via backend SMTP",
        ],
    )

    # ── 3 Backend ───────────────────────────────────────────────────────────
    doc.add_heading("3. Backend Technical Architecture", 1)

    doc.add_heading("3.1 Stack", 2)
    table(
        doc,
        ["Technology", "Role"],
        [
            ["FastAPI + Uvicorn + Pydantic v2", "HTTP API and validation"],
            ["OpenAI", "Chat, quiz MCQs, TTS, images, agent decisions"],
            ["ChromaDB + sentence-transformers", "Dense retrieval over textbooks"],
            ["rank_bm25 + CrossEncoder", "Hybrid / rerank for Maths & Science"],
            ["bcrypt", "Password hashing"],
            ["toml / python-dotenv", "Local secrets and env loading"],
        ],
    )

    doc.add_heading("3.2 Entry & Layout", 2)
    code(
        doc,
        "backend/\n"
        "├── api_server.py           # FastAPI app + all routes\n"
        "├── chat_engine.py          # Chat ↔ RAG / LLM / visual aids\n"
        "├── quiz_engine.py          # MCQ generation\n"
        "├── config/secrets.toml     # LOCAL only (gitignored)\n"
        "├── data/                   # JSON persistence\n"
        "├── quiz_data/              # Per-user quiz JSON\n"
        "├── OneSharedChromaDB/      # Vector store (ptb_textbooks)\n"
        "├── utils/\n"
        "│   ├── rag.py, llm.py, auth.py, session.py\n"
        "│   ├── email_otp.py, family_invite.py\n"
        "│   ├── media_agent.py, visual_aids.py, quiz_db.py\n"
        "│   └── secrets.py\n"
        "├── Dockerfile\n"
        "└── railway.toml",
    )

    doc.add_heading("3.3 Route Groups", 2)
    table(
        doc,
        ["Prefix", "Responsibility"],
        [
            ["/api/auth/*", "Child/parent signup, login, OTP, forgot password, /me"],
            ["/api/family/*", "Create/redeem/approve/reject/unlink invites"],
            ["/api/parent/*", "Children list, dashboard aggregates, report"],
            ["/api/users/me/*", "Profile, grade, avatar, mood, schedule, subjects"],
            ["/api/profile/learning", "Learning profile + audio prefs"],
            ["/api/chat/*", "Sessions, messages, speech, images, recap, session quiz"],
            ["/api/quiz/*", "Chapters, generate, submit, history"],
            ["/api/analytics", "Aggregated quiz analytics for student"],
            ["/api/agent/*", "Emotion, decide/run, content, MCQs, ladder, summary"],
            ["/api/contact", "Support form → SMTP + contact_messages.json"],
            ["/api/health", "Liveness for Railway"],
        ],
    )

    doc.add_heading("3.4 Authentication & OTP", 2)
    bullets(
        doc,
        [
            "Child sessions: utils/session.py → data/sessions.json (SHA-256 token)",
            "Parent sessions: data/parent_sessions.json",
            "Users: data/users.json; parents: data/parents.json; passwords via bcrypt",
            "Email OTP: utils/email_otp.py → hashed in data/email_otps.json",
            "Purposes include email_verify and password_reset (TTL ~10 minutes)",
            "Production (AUTISTUDY_ENV=production): never return plaintext OTP in API; SMTP required",
            "Dev without SMTP may expose dev_otp for local testing",
        ],
    )

    doc.add_heading("3.5 Family Linking (v6)", 2)
    bullets(
        doc,
        [
            "Invite codes FAM-XXXXX, ~48h TTL, hashed storage in data/family_invites.json",
            "Child creates invite → parent redeems → child approve/reject (app or email link)",
            "Rate limits in data/invite_rate_limits.json",
            "Max two linked parents (father/mother slots)",
        ],
    )

    doc.add_heading("3.6 RAG & Subjects", 2)
    bullets(
        doc,
        [
            "Chroma path OneSharedChromaDB (override CHROMA_DB_PATH), collection ptb_textbooks",
            "Embedder: all-MiniLM-L6-v2; reranker: ms-marco-MiniLM-L-6-v2",
            "Grades 4–5: Maths + General Science; 6–7 add Computer",
            "Maths pipeline: hybrid dense/BMASS + CrossEncoder (+ BODMAS-focused prompts)",
            "Science/Computer: hybrid + RRF + keyword gating",
            "Curriculum markdown under frontend/books_mds (Docker BOOKS_DIR=/app/books_mds)",
        ],
    )

    doc.add_heading("3.7 Quiz Engine", 2)
    bullets(
        doc,
        [
            "quiz_engine.py generates MCQs (GPT-4o-mini) with plain math notation",
            "Persistence: utils/quiz_db.py → quiz_data/<safe_email>_quiz.json",
            "Endpoints: chapters, generate, generate-from-chat, submit, history",
        ],
    )

    doc.add_heading("3.8 Persistence Map", 2)
    table(
        doc,
        ["Path", "Contents"],
        [
            ["data/users.json", "Student accounts"],
            ["data/sessions.json", "Child bearer sessions"],
            ["data/parents.json", "Parent accounts"],
            ["data/parent_sessions.json", "Parent tokens"],
            ["data/chats.json", "Chat history"],
            ["data/email_otps.json", "Hashed OTPs"],
            ["data/family_invites.json", "Family invite records"],
            ["data/contact_messages.json", "Contact form copies"],
            ["data/dashboard/<email>.json", "Mood + schedule"],
            ["data/agent_memory/", "Per-user adaptive memory"],
            ["quiz_data/", "Quiz attempts"],
            ["OneSharedChromaDB/", "Vector index"],
            ["temp_generated_images/", "Generated illustrations"],
        ],
    )

    doc.add_heading("3.9 Config & Secrets", 2)
    bullets(
        doc,
        [
            "utils/secrets.py: environment variables first, then config/secrets.toml",
            "Never commit secrets.toml; use secrets.toml.example as template",
            "Critical env: OPENAI_API_KEY, CORS_ORIGINS, AUTISTUDY_ENV, SMTP_*",
            "Optional: CHROMA_DB_PATH, CHROMA_COLLECTION, BOOKS_DIR, AUTH_DEV_OTP",
        ],
    )

    # ── 4 Deploy ────────────────────────────────────────────────────────────
    doc.add_heading("4. Deployment", 1)
    doc.add_heading("4.1 Railway (Backend)", 2)
    bullets(
        doc,
        [
            "Root Directory / Dockerfile: backend/Dockerfile built from repo root",
            "Copies backend + frontend/books_mds into image",
            "Healthcheck GET /api/health (allow long first boot for ML models)",
            "Recommend ≥2GB RAM; persist data/ and Chroma via volume if needed",
            "Set SMTP_* and OPENAI_API_KEY in Railway Variables (not git)",
        ],
    )
    doc.add_heading("4.2 Vercel (Frontend)", 2)
    bullets(
        doc,
        [
            "Root Directory = frontend",
            "Env NEXT_PUBLIC_API_URL = Railway public URL (no trailing slash)",
            "Add Vercel URL to Railway CORS_ORIGINS and redeploy backend",
        ],
    )

    # ── 5 Security ──────────────────────────────────────────────────────────
    doc.add_heading("5. Security Notes", 1)
    bullets(
        doc,
        [
            "Bearer tokens over HTTPS; separate child vs parent session stores",
            "OTP hashed at rest; password-reset grant is one-time",
            "Forgot-password responses are anti-enumeration (generic success)",
            "Camera optional with explicit consent; frames not persisted as video",
            "Production must not log or return plaintext OTPs",
        ],
    )

    # ── 6 Local Dev ─────────────────────────────────────────────────────────
    doc.add_heading("6. Local Development", 1)
    code(
        doc,
        "# From frontend/\n"
        "npm run dev:all\n"
        "# → http://localhost:3000  (web)\n"
        "# → http://127.0.0.1:8000  (API)\n\n"
        "# Optional SMTP in backend/config/secrets.toml [smtp]\n"
        "# Manual in-app: http://localhost:3000/manual",
    )

    doc.add_heading("7. Document Control", 1)
    bullets(
        doc,
        [
            "Generated by: frontend/scripts/generate_fullstack_technical_docx.py",
            "Aligned with production folders frontend/ + backend/ (v6 feature set)",
            "Companion ops: DEPLOY.md, backend/docs/EMAIL_OTP_OPS_CHECKLIST.md",
            "Older B-Form/CNIC family wording is obsolete — use FAM invite flow",
        ],
    )

    return doc


def main() -> None:
    doc = build()
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
