"""
Generate RAT evaluation Word document with Solution Correctness + Step Accuracy table.

Run from frontend/ (after rat_eval.py):
  python scripts/generate_rat_evaluation_docx.py

Output:
  docs/AutiStudy_RAT_Evaluation.docx
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

REACT_ROOT = Path(__file__).resolve().parent.parent
RESULTS_JSON = REACT_ROOT / "scripts" / "eval" / "results" / "rat_eval.json"
OUT_PATH = REACT_ROOT / "docs" / "AutiStudy_RAT_Evaluation.docx"

NAVY = RGBColor(15, 45, 74)


def style_doc(doc: Document) -> None:
    for name, size in (("Heading 1", 16), ("Heading 2", 13)):
        st = doc.styles[name]
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = NAVY


def build() -> Document:
    if not RESULTS_JSON.exists():
        raise FileNotFoundError(
            f"Missing {RESULTS_JSON}. Run: python scripts/eval/rat_eval.py"
        )

    data = json.loads(RESULTS_JSON.read_text(encoding="utf-8"))
    summary = data.get("summary_table") or []

    doc = Document()
    style_doc(doc)

    doc.add_heading("AutiStudy — RAT Evaluation Results", 0)
    doc.add_paragraph(
        "Retrieval-Augmented Thought (RAT) evaluation on procedural mathematics "
        "questions for Grades 4–7. Solution Correctness measures whether the "
        "final answer matches the gold label; Step Accuracy measures the fraction "
        "of chain-of-thought steps that pass RAT verification without correction."
    )

    doc.add_heading("Summary by Grade", 1)
    table = doc.add_table(rows=1 + len(summary), cols=3)
    table.style = "Table Grid"
    headers = ["Grade", "Solution Correctness", "Step Accuracy"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for ri, row in enumerate(summary):
        cells = table.rows[ri + 1].cells
        cells[0].text = str(row["grade"])
        cells[1].text = f"{row['solution_correctness_pct']}%"
        cells[2].text = f"{row['step_accuracy_pct']}%"

    doc.add_paragraph()

    doc.add_heading("Metric Definitions", 2)
    doc.add_paragraph(
        "Solution Correctness: percentage of test problems where RAT's generated "
        "final answer matches the expected gold answer."
    )
    doc.add_paragraph(
        "Step Accuracy: percentage of initial CoT reasoning steps that RAT "
        "verification accepts as correct (unchanged or explicitly marked correct)."
    )

    doc.add_heading("Per-Grade Detail", 1)
    for grade_key in sorted(data.get("by_grade", {}), key=int):
        block = data["by_grade"][grade_key]
        doc.add_heading(f"Grade {grade_key}", 2)
        doc.add_paragraph(
            f"Items: {block['items_evaluated']} | "
            f"Solution Correctness: {block['solution_correctness_pct']}% | "
            f"Step Accuracy: {block['step_accuracy_pct']}%"
        )
        detail = doc.add_table(
            rows=1 + len(block.get("items", [])),
            cols=5,
        )
        detail.style = "Table Grid"
        for i, h in enumerate(
            ["ID", "Query", "Expected", "Predicted", "Solution OK"]
        ):
            detail.rows[0].cells[i].text = h
        for ri, it in enumerate(block.get("items", [])):
            cells = detail.rows[ri + 1].cells
            cells[0].text = it["id"]
            cells[1].text = it["query"][:60]
            cells[2].text = it["expected_answer"]
            cells[3].text = (it.get("predicted_answer") or "")[:40]
            cells[4].text = "Yes" if it.get("solution_correct") else "No"
        doc.add_paragraph()

    note = doc.add_paragraph(
        f"Generated from {RESULTS_JSON.name} "
        f"(elapsed {data.get('elapsed_s', '?')}s total eval runtime)."
    )
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def main() -> None:
    doc = build()
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
