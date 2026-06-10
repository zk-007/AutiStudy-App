"""
Generate Section 4 (System Design & Architecture) Word document for the AutiStudy
research paper, following AutiStudy_Final_Research_Guide.pdf structure.

Run: python scripts/generate_section4_architecture_docx.py
Output: docs/AutiStudy_Section4_System_Architecture.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT_PATH = Path(__file__).resolve().parent.parent / "docs" / "AutiStudy_Section4_System_Architecture.docx"


def style_doc(doc: Document) -> None:
    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(15, 45, 74)
    h2 = doc.styles["Heading 2"]
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(37, 99, 235)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            t.rows[ri + 1].cells[ci].text = cell
    doc.add_paragraph()


def build() -> Document:
    doc = Document()
    style_doc(doc)

    # Title block
    title = doc.add_heading("IV. SYSTEM DESIGN AND ARCHITECTURE", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        "AutiStudy: A Pilot Evaluation of an AI-Powered Virtual Shadow Teacher "
        "for Autistic Students in Pakistan"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph(
        "Section 4 — System Design & Architecture | Author: Zoha Khalid | "
        "Target length: 1,000–1,200 words (Q2 journal standard)"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 4.1
    doc.add_heading("4.1 System Overview", 1)
    doc.add_paragraph(
        "AutiStudy is implemented as a four-layer architecture that separates user interaction, "
        "API orchestration, curriculum-grounded intelligence, and adaptive tutoring policy. "
        "As illustrated in Fig. 1, Layer 1 is the Next.js 14 web frontend (student and parent "
        "interfaces, bilingual RTL layout, in-browser visual aids, and client-side computer vision). "
        "Layer 2 is a single FastAPI backend (api_server.py) exposed over HTTP on port 8000, "
        "which handles authentication, chat persistence, quiz analytics, and multimodal content "
        "generation. Layer 3 is the AI core: hybrid retrieval (ChromaDB dense search, BM25 sparse "
        "retrieval, reciprocal rank fusion, and cross-encoder re-ranking), GPT-4o Mini generation, "
        "OpenAI TTS, and DALL·E-class image generation. Layer 4 is the agentic layer: nine "
        "coordinated logical agents (Supervisor, Content, RAG, Critic, Media, Personalisation, "
        "Sensory, Recap, and Progress) that together implement real-time adaptation, safety checks, "
        "and long-term student memory."
    )
    doc.add_paragraph(
        "The frontend and backend live in one monorepo (AutiStudy-App: frontend/ + backend/) "
        "and communicate via a typed REST client with Bearer-token authentication. "
        "Curriculum markdown derived from Pakistan Single National Curriculum (SNC) textbooks "
        "(Grades 4–7: Mathematics, General Science, Computer Science) is ingested into ChromaDB "
        "after LLaMA Parse extraction. Production tutoring combines server-side RAG answers with "
        "browser-side hybrid emotion fusion (MediaPipe + face-api.js) and a gated comprehension-check "
        "protocol (“Did you get it?”) that escalates through flowcharts, read-aloud, images, MCQs, "
        "and breathing exercises when understanding is not confirmed."
    )
    doc.add_paragraph(
        "[Fig. 1 — Four-layer system architecture: Frontend → FastAPI → AI Core → Agentic Layer. "
        "Export as vector SVG or ≥300 DPI PNG from draw.io/Figma.]"
    )

    # 4.2
    doc.add_heading("4.2 Frontend Architecture", 1)
    doc.add_paragraph(
        "The student-facing application is autistudy-web (Next.js 14.2.35, App Router, React 18, "
        "TypeScript 5). Table I summarises the stack and design rationale for an autism-friendly, "
        "low-distraction learning environment."
    )
    add_table(
        doc,
        ["Layer", "Technology", "Version", "Justification"],
        [
            ["Framework", "Next.js (App Router)", "14.2.35", "Fast initial load; file-based routing for student/parent roles"],
            ["UI", "React", "18", "Component model; concurrent rendering for smooth UI"],
            ["Language", "TypeScript", "5", "Compile-time safety for child-facing flows"],
            ["Styling", "Tailwind CSS", "3.4.1", "Utility-first CSS; CSS variables for real-time sensory themes"],
            ["Animation", "Framer Motion", "12.38", "Transitions disable globally via reduce-motion (sensory-averse users)"],
            ["Scroll", "Lenis", "1.3", "Smooth scroll reduces jarring page jumps"],
            ["Landing", "GSAP", "3.15", "GPU landing animations isolated from tutoring UI"],
            ["Math", "react-markdown + KaTeX", "—", "Server-rendered math; normalizeMath() for delimiters"],
            ["Icons", "lucide-react", "1.8", "Accessible icons with ARIA labels"],
            ["Fonts", "Quicksand + Inter", "Google Fonts", "Rounded forms aid dyslexia/ASD readability"],
            ["E2E", "Playwright", "1.59", "Automated login, chat, quiz, parent flows"],
            ["API client", "lib/api/client.ts", "—", "Typed fetch; dual tokens (student/parent)"],
        ],
    )
    doc.add_paragraph("Table I — Frontend technology stack.")
    doc.add_paragraph(
        "Key routes include / (landing), /login and /signup (dual role), /dashboard (stats and "
        "subject cards), /chat (AI tutor with URL state machine: ?subject → ?session), /quiz "
        "(phase-based MCQs), /analytics (seven-day activity), /parent/dashboard and /parent/report "
        "(child metrics and LLM-generated JSON reports). The chat page orchestrates two coordinated "
        "subsystems: useAdaptiveTutorAgent (hybrid webcam CV, engagement analyser, preferred answer "
        "format hints) and useComprehensionFlow (popup-gated adaptation ladder with scroll/TTS/image "
        "timing gates). Raw webcam frames never leave the browser; only derived emotion probabilities "
        "influence tutoring policy, preserving privacy for minors."
    )
    doc.add_paragraph(
        "Eleven in-browser visual aid types (EmojiCounting, MathStepCard, FactorTree, FractionBar, "
        "NumberLine, BarChart, PercentageBar, TimesTable, Geometry, Ratio, and DALL·E concept images) "
        "render client-side from structured JSON returned by the backend, avoiding extra network "
        "latency after the initial response. LocaleProvider enables English/Urdu switching with RTL "
        "layout; SettingsContext persists font size, high contrast, focus mode, and animation toggles "
        "to localStorage (WCAG-aligned: resize text SC 1.4.4, contrast SC 1.4.3, animation SC 2.3.3)."
    )

    # 4.3
    doc.add_heading("4.3 Backend Architecture", 1)
    doc.add_paragraph(
        "The backend uses FastAPI only (not Flask): a single api_server.py process with uvicorn, "
        "CORS for localhost:3000/3001, and a 16-worker ThreadPoolExecutor (run_in_thread) so "
        "blocking ChromaDB, EasyOCR, and bcrypt operations do not stall the async event loop. "
        "Authentication employs 32-character hex Bearer tokens (SHA-256 derived, seven-day TTL) with "
        "bcrypt password hashing (12 rounds; legacy SHA-256 migration on login). Persistence is "
        "file-based JSON under data/ (users, sessions, chats, parents, agent_memory) plus "
        "quiz_data/ per student—appropriate for a research prototype with a single Next.js UI."
    )
    add_table(
        doc,
        ["Group", "Representative endpoints"],
        [
            ["Public", "GET /, /api/health, /api/chat/config"],
            ["Student auth", "POST /api/auth/register, /login, /logout; GET /api/auth/me; POST /api/auth/child/signup (multipart B-Form)"],
            ["Parent auth", "POST /api/auth/parent/signup, /login; GET /api/parent/dashboard, /report"],
            ["Dashboard", "GET /api/users/me/stats, /subjects, /recent-chats, /recent-quizzes"],
            ["Chat", "GET/POST /api/chat/sessions; POST .../messages, /image, /speech, /quiz"],
            ["Quiz & analytics", "GET /api/quiz/chapters; POST /generate, /submit; GET /api/analytics"],
            ["Teaching agent", "POST /api/agent/generate-content, /step-mcqs; GET /api/agent/memory; POST /session-summary"],
            ["Legacy CV API", "POST /api/agent/analyze-emotion (GPT-4o Mini Vision on JPEG frames)"],
        ],
    )
    doc.add_paragraph("Table II — REST API endpoint groups.")
    doc.add_paragraph(
        "chat_engine.py bridges the API to utils/llm.py and utils/rag.py. Each chat message triggers "
        "intent-aware retrieval, prompt construction with grade/subject context, and gpt-4o-mini "
        "generation (max ~600 tokens, temperature 0.7). Visual aids are routed deterministically by "
        "utils/visual_aids.py (regex priority, no LLM) before optional DALL·E fallback. Generated "
        "images are cached under temp_generated_images/ and served at /api/generated-images/."
    )

    # 4.4
    doc.add_heading("4.4 Data Collection and LLaMA Parse Ingestion", 1)
    doc.add_paragraph(
        "Curriculum data comprise Pakistan SNC textbook PDFs for Grades 4–7 (Mathematics, General "
        "Science, Computer Science). Initial OCR (EasyOCR + OpenCV) achieved only ~45% usable text "
        "quality—mathematical notation, Urdu script, and tables were frequently malformed. LLaMA Parse "
        "(LlamaIndex document parser) raised extraction quality to ~98%, preserving LaTeX, headings, "
        "tables, and Urdu structure. Chunks are semantic segments with metadata: Chapter Name, Section "
        "Title, Page Number, and Block Type (Definition, Procedure, Example, Practice, Explanation). "
        "Procedure and Explanation blocks receive higher retrieval weights because they align with "
        "procedural tutoring queries."
    )
    doc.add_paragraph(
        "Because Urdu subword tokenization consumes more OpenAI tokens per word than Latin text, "
        "chunk sizes should be reported in both characters and estimated tokens (e.g., ~800 characters "
        "≈ ~200 English tokens vs. ~350 tokens for equivalent Urdu content). This justifies the k=6 "
        "context limit at generation time (Salemi & Zamani). Table III (Dataset Statistics) should list "
        "books ingested, total pages/chunks, vocabulary size, and block-type distribution once computed "
        "from the ChromaDB collection ptb_textbooks."
    )
    doc.add_paragraph(
        "Known limitation: math_parse_7.md exists on disk but is not registered in BOOK_MAP, so "
        "Grade 7 Mathematics quiz chapters are unavailable in the current build."
    )

    # 4.5
    doc.add_heading("4.5 Hybrid RAG Pipeline (Science and Computer Science)", 1)
    doc.add_paragraph(
        "Science and Computer Science share a hybrid retrieval pipeline (see Fig. 2). Step 1 applies "
        "rule-based intent detection (definition, example, procedure, problem-solving). Step 2 runs "
        "parallel dense semantic search (ChromaDB, all-MiniLM-L6-v2) and BM25 sparse retrieval—weighted "
        "65% dense / 35% BM25 for Computer Science; Science uses reciprocal rank fusion (RRF, k=60) "
        "after parallel retrieval. Step 3 merges ranked lists via RRF. Step 4 applies keyword gating "
        "and metadata filtering so block types match detected intent. Step 5 re-ranks the top 40 "
        "candidates with cross-encoder/ms-marco-MiniLM-L-6-v2 and selects top k=6 documents. Step 6 "
        "injects these chunks into GPT-4o Mini for a curriculum-aligned answer. Offline evaluation "
        "reports Retrieval Precision@5 of 87% and Answer Relevance of 91% versus OCR-based baselines "
        "(+45% and +36% respectively)."
    )

    # 4.6
    doc.add_heading("4.6 RAT Pipeline (Mathematics)", 1)
    doc.add_paragraph(
        "Mathematics uses Retrieval-Augmented Thinking (RAT) in evaluation: GPT-4o Mini first decomposes "
        "a problem via chain-of-thought into sub-questions; each sub-query searches ChromaDB independently; "
        "retrieved Procedure blocks verify intermediate steps; a step-by-step explanation is synthesised. "
        "This design achieves MRR = 1.00 for procedural queries versus 0.255 for dense-only RAG in offline "
        "tests—because sub-queries align with Procedure metadata rather than concept definitions."
    )
    doc.add_paragraph(
        "Important implementation note: RAT is currently disabled in the production API (use_rat flag) "
        "due to latency from multiple cross-encoder passes per sub-query. Live maths tutoring uses the "
        "optimised hybrid RAG path (65% dense + 35% BM25) until RAT is parallelised or cached. Fig. 3 "
        "should depict the RAT pipeline for the paper; evaluation metrics come from offline benchmarks."
    )

    # 4.7
    doc.add_heading("4.7 Nine-Agent Agentic Architecture", 1)
    doc.add_paragraph(
        "AutiStudy coordinates nine logical agents (Fig. 4) to mirror a human shadow teacher: routing, "
        "retrieval, critique, media, personalisation, sensory UI, session recap, and progress reporting. "
        "At the design level, LangGraph StateGraph orchestrates agent hand-offs; in the deployed prototype, "
        "equivalent behaviour is distributed across FastAPI modules and React state machines."
    )
    add_table(
        doc,
        ["Agent", "Role", "Implementation"],
        [
            ["Supervisor", "Routes queries to RAG or RAT; coordinates adaptation", "api_server + chat_engine orchestration"],
            ["Content", "Intent detection; builds retrieval plan", "Rule-based classifier + gpt-4o-mini"],
            ["RAG", "Hybrid retrieval and re-ranking", "utils/rag.py — ChromaDB, BM25, MS-Marco"],
            ["Critic", "Accuracy, clarity, age-appropriateness (CoT)", "Prompt rules + relevance thresholds"],
            ["Media", "Images (gpt-image chain) and TTS (tts-1)", "utils/llm.py, visual_aids.py"],
            ["Personalisation", "Difficulty from quiz history", "agent_memory JSON + analytics"],
            ["Sensory", "Font, contrast, motion, sound toggles", "SettingsContext + CSS variables (React)"],
            ["Recap", "Session continuity summaries", "gpt-4o-mini via session-summary endpoint"],
            ["Progress", "Long-term memory; parent dashboard", "agent_memory/{md5}.json, parent APIs"],
        ],
    )
    doc.add_paragraph("Table III — Nine-agent roles and technologies.")
    doc.add_paragraph(
        "The browser adaptive tutor extends the Sensory and Supervisor roles: HybridEmotionEngine fuses "
        "MediaPipe blendshapes (~15 FPS) with face-api.js CNN expressions (~2 FPS), smoothed over 3–8 s "
        "buffers. ComprehensionStateMachine and TutorComprehensionFlow implement the five-round ladder "
        "(flowchart → read-aloud → image → recall/teaching MCQs → breathing). POST /api/agent/generate-content "
        "produces adaptation text; POST /api/agent/step-mcqs supports default, teaching, and appreciation "
        "modes. A legacy ReAct Media Agent (POST /api/agent/run, GPT-4o function calling, max three "
        "iterations) remains for full tool orchestration but is separate from the primary popup-gated flow."
    )

    # 4.8
    doc.add_heading("4.8 Security and Prompt-Injection Defence", 1)
    doc.add_paragraph(
        "Because AutiStudy serves minors, defence-in-depth is applied. The Critic agent concept evaluates "
        "answers for curriculum alignment and age-appropriateness before display. Content-agent intent "
        "classification rejects out-of-scope queries (non Maths/Science/Computer Science) with a boundary "
        "message. Subject keyword gating filters retrieved chunks from unrelated subjects. System prompts "
        "are fixed at the API layer; user text is passed only in user-role messages to reduce role-confusion "
        "attacks. RAG relevance scoring (threshold 0.35 with subject keywords, 0.55 without) blocks "
        "off-textbook queries (~85% detection vs. ~20% baseline in internal tests). Student-facing agents "
        "have no web browsing, code execution, or arbitrary external APIs—attack surface is limited to "
        "the curriculum knowledge base. Camera consent is stored locally (autistudy_camera_consent); CNIC "
        "data from B-Form OCR is used for registration verification and not retained in plaintext after "
        "account creation. The debug endpoint /api/debug/openai-ping must be removed before production."
    )

    # Figures checklist
    doc.add_heading("Figures and Tables Checklist (Section IV)", 1)
    add_table(
        doc,
        ["Item", "Description", "Status"],
        [
            ["Fig. 1", "Four-layer architecture (Frontend → FastAPI → AI Core → Agents)", "To export ≥300 DPI"],
            ["Fig. 2", "Hybrid RAG pipeline (Science/CS)", "To export ≥300 DPI"],
            ["Fig. 3", "RAT pipeline (Mathematics)", "To export ≥300 DPI"],
            ["Fig. 4", "Nine-agent architecture", "To export ≥300 DPI"],
            ["Table I", "Frontend stack", "Included above"],
            ["Table II", "API endpoint groups", "Included above"],
            ["Table III", "Nine-agent mapping", "Included above"],
            ["Table (paper)", "Dataset statistics (Section 6.2)", "Fill from ChromaDB"],
        ],
    )

    doc.add_paragraph()
    note = doc.add_paragraph(
        "Document generated from AutiStudy_Final_Research_Guide.pdf (Section 4) and aligned with "
        "AutiStudy-App monorepo implementation (v4.1 adaptive-agent, June 2026)."
    )
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def main() -> None:
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    build().save(str(OUT_PATH))
    print(f"Created: {OUT_PATH}")


if __name__ == "__main__":
    main()
