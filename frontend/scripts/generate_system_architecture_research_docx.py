"""
Generate SYSTEM_ARCHITECTURE_RESEARCH.docx from the markdown source.

Run from frontend/:
  python scripts/generate_system_architecture_research_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

MD_PATH = Path(__file__).resolve().parent.parent.parent / "backend" / "docs" / "SYSTEM_ARCHITECTURE_RESEARCH.md"
OUT_PATH = MD_PATH.with_suffix(".docx")


def style_doc(doc: Document) -> None:
    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(15, 45, 74)
    h2 = doc.styles["Heading 2"]
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(37, 99, 235)
    h3 = doc.styles["Heading 3"]
    h3.font.size = Pt(11)
    h3.font.bold = True


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Pt(12)


def parse_table_rows(block: list[str]) -> tuple[list[str], list[list[str]]] | None:
    if len(block) < 2:
        return None
    rows = []
    for line in block:
        if not line.strip().startswith("|"):
            return None
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) < 2:
        return None
    # skip markdown separator row |---|---|
    if all(re.match(r"^:?-+:?$", c.replace(" ", "")) or c == "" for c in rows[1]):
        data_rows = rows[2:]
        header = rows[0]
    else:
        header = rows[0]
        data_rows = rows[1:]
    return header, data_rows


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci in range(len(headers)):
            text = row[ci] if ci < len(row) else ""
            t.rows[ri + 1].cells[ci].text = text
    doc.add_paragraph()


def add_paragraph_text(doc: Document, text: str) -> None:
    text = text.strip()
    if not text:
        return
    # strip markdown bold/italic/code lightly
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    if text.startswith("> "):
        p = doc.add_paragraph(text[2:])
        p.paragraph_format.left_indent = Pt(18)
        for run in p.runs:
            run.italic = True
    else:
        doc.add_paragraph(text)


def convert_md_to_docx(md_path: Path, out_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    style_doc(doc)

    i = 0
    code_buf: list[str] = []
    in_code = False
    table_buf: list[str] = []

    def flush_table() -> None:
        nonlocal table_buf
        if not table_buf:
            return
        parsed = parse_table_rows(table_buf)
        if parsed:
            add_table(doc, parsed[0], parsed[1])
        table_buf = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                flush_table()
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if line.strip().startswith("|"):
            table_buf.append(line)
            i += 1
            continue
        flush_table()

        if line.startswith("# ") and not line.startswith("## "):
            doc.add_heading(line[2:].strip(), 0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), 1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), 2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), 3)
        elif line.strip() == "---":
            doc.add_paragraph()
        elif line.strip():
            add_paragraph_text(doc, line)
        i += 1

    flush_table()
    if in_code and code_buf:
        add_code_block(doc, code_buf)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    if not MD_PATH.exists():
        raise SystemExit(f"Missing source: {MD_PATH}")
    convert_md_to_docx(MD_PATH, OUT_PATH)
