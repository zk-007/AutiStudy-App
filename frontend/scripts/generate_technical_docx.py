"""
Generate backend.docx and frontend.docx technical documentation for AutiStudy.
Run: python scripts/generate_technical_docx.py
"""
from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = Path(__file__).resolve().parent.parent


def style_doc(doc: Document) -> None:
    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(37, 99, 235)
    h2 = doc.styles["Heading 2"]
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(15, 45, 74)


def title_page(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_heading(title, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_paragraph(subtitle)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    meta = doc.add_paragraph("AutiStudy — Technical Documentation")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            t.rows[ri + 1].cells[ci].text = cell
    doc.add_paragraph()


def code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.style = "Intense Quote"


def build_backend() -> Document:
    doc = Document()
    style_doc(doc)
    title_page(
        doc,
        "AutiStudy Backend",
        "Technical documentation — all backend work completed before the Media Agent",
    )

    # 1. Overview
    doc.add_heading("1. Executive Summary", 1)
    doc.add_paragraph(
        "The AutiStudy backend is a Python stack centered on FastAPI (api_server.py) that exposes "
        "a REST API for the Next.js React frontend. It reuses the same JSON persistence layer and "
        "AI modules originally built for the Streamlit app, ensuring both UIs stay in sync. "
        "Core capabilities include JWT-style bearer auth, RAG-grounded tutoring, deterministic "
        "visual-aid routing, quiz generation, analytics, parent dashboards, B-Form OCR signup, "
        "emotion-aware teaching modality, and OpenAI TTS."
    )
    doc.add_paragraph(
        "Scope note: This document covers all backend work EXCEPT the Media Agent "
        "(utils/media_agent.py, POST /api/agent/run). The Media Agent is a separate ReAct "
        "orchestrator added later; only a brief exclusion boundary is listed in Section 14."
    )

    # 2. Architecture
    doc.add_heading("2. System Architecture", 1)
    doc.add_heading("2.1 Repository Layout", 2)
    code_block(
        doc,
        "AutiStudy/\n"
        "├── api_server.py          # FastAPI REST API (primary entry for React)\n"
        "├── app.py                 # Streamlit multi-page app (legacy UI)\n"
        "├── chat_engine.py         # Bridge: API ↔ RAG/LLM/visual aids\n"
        "├── quiz_engine.py         # GPT quiz generation\n"
        "├── requirements.txt\n"
        "├── Dockerfile             # Streamlit deployment (port 7860)\n"
        "├── scripts/run_api_detached.py\n"
        "├── data/                  # JSON persistence\n"
        "├── quiz_data/             # Per-user quiz attempts\n"
        "├── OneSharedChromaDB/     # ChromaDB vector store\n"
        "├── temp_generated_images/ # Cached DALL·E outputs\n"
        "├── utils/                 # Core modules\n"
        "└── views/                 # Streamlit pages\n\n"
        "AutiStudy-React/books_mds/  # Curriculum markdown (sibling repo)",
    )
    doc.add_heading("2.2 Request Flow", 2)
    bullets(
        doc,
        [
            "React frontend → HTTP → FastAPI api_server.py",
            "Auth: Bearer token from utils/session.py (student) or parent_sessions.json (parent)",
            "Chat: api_server → chat_engine → utils/llm.py + utils/rag.py",
            "Visual aids: chat_engine → utils/visual_aids.py (deterministic router)",
            "Quiz: api_server → quiz_engine.py + utils/book_parser.py",
            "Blocking I/O runs in ThreadPoolExecutor (16 workers) via run_in_thread()",
        ],
    )

    # 3. FastAPI Server
    doc.add_heading("3. FastAPI Server (api_server.py)", 1)
    table(
        doc,
        ["Property", "Value"],
        [
            ["Framework", "FastAPI 0.110+"],
            ["ASGI", "uvicorn"],
            ["Default URL", "http://127.0.0.1:8000"],
            ["Version", "0.1.0"],
            ["CORS origins", "localhost:3000, 3001, 127.0.0.1:3000"],
            ["Static mount", "/api/generated-images → temp_generated_images/"],
            ["Thread pool", "16 workers for blocking calls"],
        ],
    )
    doc.add_heading("3.1 Startup Behavior", 2)
    bullets(
        doc,
        [
            "Background warmup thread: loads OpenAI key, preloads RAG (embedder, reranker, ChromaDB)",
            "OpenAI HTTPS keep-alive ping every 60 seconds (Windows TLS handshake mitigation)",
            "UTF-8 forced on stdout/stderr on Windows",
            "Health and config endpoints available while models load",
        ],
    )

    # 4. API Endpoints
    doc.add_heading("4. REST API Endpoints", 1)

    doc.add_heading("4.1 Public", 2)
    table(
        doc,
        ["Method", "Path", "Description"],
        [
            ["GET", "/", "API name, status, version"],
            ["GET", "/api/health", "{ok: true}"],
            ["GET", "/api/chat/config", "tutor_configured, rag_available, images_available, speech_available"],
            ["GET", "/api/debug/openai-ping", "OpenAI connectivity test (remove in production)"],
        ],
    )

    doc.add_heading("4.2 Student Authentication", 2)
    table(
        doc,
        ["Method", "Path", "Body / Notes"],
        [
            ["POST", "/api/auth/register", "name, email, password, grade (4-7), role"],
            ["POST", "/api/auth/login", "email, password → {token, user}"],
            ["POST", "/api/auth/logout", "Bearer token"],
            ["GET", "/api/auth/me", "Current user profile"],
            ["POST", "/api/auth/child/signup", "multipart: name, email, password, grade, cnic, bform file"],
            ["POST", "/api/users/me/password", "current_password, new_password"],
        ],
    )
    doc.add_paragraph(
        "Password validation: minimum 8 characters, 1 uppercase, 1 lowercase, 1 digit, 1 special character."
    )

    doc.add_heading("4.3 Parent Authentication", 2)
    table(
        doc,
        ["Method", "Path", "Description"],
        [
            ["POST", "/api/auth/parent/signup", "Parent + child CNIC verification"],
            ["POST", "/api/auth/parent/login", "Parent login"],
            ["POST", "/api/auth/parent/logout", "Invalidate parent token"],
            ["GET", "/api/auth/parent/me", "Parent profile"],
            ["GET", "/api/parent/dashboard", "Child analytics dashboard"],
            ["GET", "/api/parent/report", "LLM-generated JSON progress report"],
        ],
    )

    doc.add_heading("4.4 User Dashboard", 2)
    table(
        doc,
        ["Method", "Path", "Returns"],
        [
            ["GET", "/api/users/me/stats", "stars, streak, quiz totals, accuracy, daily_activity"],
            ["GET", "/api/users/me/subjects", "Grade subjects + last_studied"],
            ["GET", "/api/users/me/recent-chats", "Up to 5 recent chat sessions"],
            ["GET", "/api/users/me/recent-quizzes", "Up to 5 recent quiz attempts"],
        ],
    )

    doc.add_heading("4.5 Chat", 2)
    table(
        doc,
        ["Method", "Path", "Description"],
        [
            ["GET", "/api/chat/sessions", "List sessions (newest first)"],
            ["POST", "/api/chat/sessions", "Create: subject, language (en|ur)"],
            ["GET", "/api/chat/sessions/{id}", "Full session + messages"],
            ["DELETE", "/api/chat/sessions/{id}", "Delete session"],
            ["POST", "/api/chat/sessions/{id}/messages", "Send message (max 4000 chars)"],
            ["POST", "/api/chat/sessions/{id}/image", "Generate visual aid for last question"],
            ["POST", "/api/chat/speech", "TTS: text, language → base64 MP3"],
            ["POST", "/api/chat/sessions/{id}/quiz", "Quiz from chat transcript"],
        ],
    )

    doc.add_heading("4.6 Quiz & Analytics", 2)
    table(
        doc,
        ["Method", "Path", "Description"],
        [
            ["GET", "/api/quiz/chapters?subject=", "Textbook chapters for grade+subject"],
            ["POST", "/api/quiz/generate", "Generate quiz (optional topic/chapter)"],
            ["POST", "/api/quiz/generate-from-chat", "Quiz from most recent chat"],
            ["POST", "/api/quiz/submit", "Score, persist, award stars"],
            ["GET", "/api/quiz/history", "Last 20 attempts"],
            ["GET", "/api/analytics", "Full student analytics"],
        ],
    )

    doc.add_heading("4.7 Teaching Agent (non–Media Agent)", 2)
    table(
        doc,
        ["Method", "Path", "Description"],
        [
            ["POST", "/api/agent/analyze-emotion", "Webcam JPEG → emotion + modality decision"],
            ["GET", "/api/agent/state/{session_id}", "Current modality state"],
            ["POST", "/api/agent/reset", "Reset agent state for session"],
            ["POST", "/api/agent/force-modality", "Manual modality override"],
            ["POST", "/api/agent/session-summary", "Persist session to long-term memory"],
            ["GET", "/api/agent/memory", "Memory summary for student"],
        ],
    )
    doc.add_paragraph(
        "Excluded from this document: POST /api/agent/run — Media Agent ReAct loop (GPT-4o function calling)."
    )

    # 5. Auth
    doc.add_heading("5. Authentication & Authorization", 1)
    doc.add_heading("5.1 Student Auth (utils/auth.py, utils/session.py)", 2)
    bullets(
        doc,
        [
            "Password hashing: bcrypt (12 rounds); legacy SHA-256 supported with migration on login",
            "Session token: 32-char hex (SHA-256 of email:timestamp)",
            "Session TTL: 7 days; expired tokens deleted on access",
            "Bearer header: Authorization: Bearer <token>",
            "User store: data/users.json keyed by email",
        ],
    )
    doc.add_heading("5.2 User Record Schema", 2)
    code_block(
        doc,
        '{\n  "name": "Student",\n  "email": "student@example.com",\n'
        '  "password": "<bcrypt hash>",\n  "role": "student",\n  "grade": 4,\n'
        '  "stars": 0, "badges": [], "progress": {},\n'
        '  "cnic": "12345-6789012-3",\n  "father_cnic": "...", "mother_cnic": "...",\n'
        '  "bform_verified": true\n}',
    )
    doc.add_heading("5.3 Child B-Form Signup", 2)
    bullets(
        doc,
        [
            "POST /api/auth/child/signup — multipart form with B-Form image",
            "CNIC normalized to XXXXX-XXXXXXX-X (13 digits)",
            "utils/ocr.py: EasyOCR + OpenCV preprocessing",
            "Validates NADRA CRC/B-Form document; extracts parent CNICs by frequency",
            "Requires easyocr and opencv-python (not in requirements.txt — install separately)",
        ],
    )
    doc.add_heading("5.4 Parent Auth (utils/parent_db.py)", 2)
    bullets(
        doc,
        [
            "Parent store: data/parents.json",
            "Parent sessions: data/parent_sessions.json (separate from student sessions)",
            "Signup verifies parent CNIC matches child father_cnic or mother_cnic from B-Form",
            "Child linked by name (+ optional child CNIC)",
        ],
    )

    # 6. Storage
    doc.add_heading("6. Data Storage (JSON Files)", 1)
    doc.add_paragraph("No SQL database. Streamlit and FastAPI share the same files.")
    table(
        doc,
        ["Store", "Path", "Purpose"],
        [
            ["Users", "data/users.json", "Student accounts"],
            ["Sessions", "data/sessions.json", "Student bearer tokens"],
            ["Parent sessions", "data/parent_sessions.json", "Parent tokens"],
            ["Parents", "data/parents.json", "Parent accounts"],
            ["Chats", "data/chats.json", "{email: [sessions]}"],
            ["Quiz", "quiz_data/{email_safe}_quiz.json", "Per-user attempts"],
            ["Agent memory", "data/agent_memory/{md5(email)}.json", "Cross-session learning"],
            ["Vectors", "OneSharedChromaDB/", "Chroma collection ptb_textbooks"],
            ["Images", "temp_generated_images/", "Cached generated images"],
        ],
    )
    doc.add_heading("6.1 Chat Session Schema", 2)
    code_block(
        doc,
        '{\n  "id": "a1b2c3d4",\n  "grade": 4, "subject": "Maths",\n'
        '  "language": "en", "title": "Maths - Mar 15",\n'
        '  "messages": [\n    {\n      "role": "user|assistant",\n'
        '      "content": "...", "timestamp": "ISO8601",\n'
        '      "image_url": null,\n      "math_steps": {...}, "emoji_counting": {...},\n'
        '      "factor_tree": {...}, ...\n    }\n  ]\n}',
    )
    doc.add_heading("6.2 Stars Award (quiz submit)", 2)
    doc.add_paragraph("90%+ → 5 stars, 80%+ → 4, 70%+ → 3, 60%+ → 2, else 1.")

    # 7. Chat Engine
    doc.add_heading("7. Chat Engine (chat_engine.py)", 1)
    bullets(
        doc,
        [
            "Bridge between FastAPI and utils/llm.py + utils/rag.py",
            "Boot: load OpenAI key → set OPENAI_API_KEY → absolute CHROMA path → lazy import llm",
            "is_configured() — OpenAI key present",
            "rag_status() — Chroma, embedder, reranker diagnostics",
            "generate_reply() — RAG-grounded tutor text",
            "generate_visual_aid() — routes to visual aid tracks",
            "generate_speech_b64() — TTS wrapper",
        ],
    )

    # 8. LLM
    doc.add_heading("8. LLM Layer (utils/llm.py)", 1)
    table(
        doc,
        ["Setting", "Value"],
        [
            ["Chat model", "gpt-4o-mini"],
            ["Max tokens", "600"],
            ["Temperature", "0.7"],
            ["Image models", "gpt-image-1.5 → gpt-image-1 → dall-e-3 (fallback chain)"],
            ["Vision verification", "gpt-4o (math image count check, up to 3 retries)"],
            ["TTS model", "tts-1"],
            ["TTS voices", "alloy (en), nova (ur)"],
            ["HTTP timeout", "60s"],
            ["Client", "Singleton httpx with keepalive"],
        ],
    )
    doc.add_heading("8.1 System Prompts", 2)
    bullets(
        doc,
        [
            "SYSTEM_PROMPT_EN / SYSTEM_PROMPT_UR — autism-friendly, grade/subject context",
            "LaTeX formatting rules for fractions and algebra",
            "Division/multiplication method rules aligned with Pakistani curriculum",
            "Emoji counting guidance for small whole numbers",
        ],
    )
    doc.add_heading("8.2 RAG Integration in generate_response()", 2)
    bullets(
        doc,
        [
            "query_knowledge_base(user_message, grade, subject)",
            "Append retrieved context to system prompt when relevant",
            "Inject 'not in textbook' or 'wrong subject' when relevance fails",
            "Last 10 chat messages included in context",
        ],
    )
    doc.add_heading("8.3 Image Pipeline", 2)
    bullets(
        doc,
        [
            "enhance_image_prompt() — GPT plans JSON visual template",
            "Generate via OpenAI image API",
            "Countable math: GPT-4o Vision verifies object counts",
            "should_auto_generate_image() always returns False — images require explicit /image call",
        ],
    )

    # 9. RAG
    doc.add_heading("9. RAG Layer (utils/rag.py)", 1)
    table(
        doc,
        ["Config", "Default"],
        [
            ["CHROMA_DB_PATH", "OneSharedChromaDB"],
            ["CHROMA_COLLECTION", "ptb_textbooks"],
            ["Embedder", "sentence-transformers/all-MiniLM-L6-v2"],
            ["Reranker", "cross-encoder/ms-marco-MiniLM-L-6-v2"],
            ["Math doc ID", "MATH4"],
            ["Science doc ID", "GS4"],
            ["Computer doc ID", "CS6"],
        ],
    )
    doc.add_heading("9.1 Subject-Specific Pipelines", 2)
    table(
        doc,
        ["Subject", "Algorithm"],
        [
            ["Maths", "Query rewrite → intent → 65% dense + 35% BM25 → cross-encoder rerank → block filter. RAT disabled in API for latency."],
            ["Computer", "Dense + BM25 → RRF fusion (k=60) → block weights → keyword gate → intent filter"],
            ["General Science", "Same hybrid+RRF as Computer"],
        ],
    )
    doc.add_heading("9.2 Relevance Scoring", 2)
    bullets(
        doc,
        [
            "compute_relevance_score() — average of top 3 hit scores",
            "Threshold: 0.35 with subject keywords; 0.55 without",
            "Returns: documents, relevance_score, is_relevant, query_related_to_subject",
        ],
    )

    # 10. Visual Aids
    doc.add_heading("10. Visual Aid Generation", 1)
    doc.add_paragraph(
        "Router: utils/visual_aids.classify_visual_request() — deterministic regex priority (no LLM). "
        "Orchestrator: chat_engine.generate_visual_aid() → POST /api/chat/sessions/{id}/image."
    )
    table(
        doc,
        ["Priority", "Track", "Trigger", "Response kind"],
        [
            ["1", "countable", "Small whole-number +/−/× (≤10)", "emoji_counting"],
            ["2", "factor_tree", "HCF, LCM, prime factors", "factor_tree"],
            ["3", "number_line", "Integers, negatives", "number_line"],
            ["4", "bar_chart", "Chart/graph keywords", "bar_chart"],
            ["5", "percentage_bar", "Percent questions", "percentage_bar"],
            ["6", "times_table", "Times table keywords", "times_table"],
            ["7", "geometry", "Shapes, perimeter, area", "geometry"],
            ["8", "ratio", "Ratio/proportion", "ratio"],
            ["9", "fraction_bar", "Fraction + visual keyword", "fraction_bar"],
            ["10", "symbolic", "Fractions, algebra, long division", "math_steps"],
            ["11", "concept", "Default", "image (DALL·E)"],
        ],
    )
    doc.add_heading("10.1 Payload Types", 2)
    bullets(
        doc,
        [
            "emoji_counting: n1, n2, op, result, emoji, labels, title",
            "math_steps: title, steps[{caption, latex}], final_answer (GPT JSON)",
            "factor_tree: ladders, prime_factors, hcf, lcm, common_factors",
            "fraction_bar, number_line, bar_chart, percentage_bar, times_table, geometry, ratio",
            "image: DALL·E URL or /api/generated-images/... local path",
        ],
    )

    # 11. Quiz
    doc.add_heading("11. Quiz System", 1)
    doc.add_heading("11.1 quiz_engine.py", 2)
    bullets(
        doc,
        [
            "generate_quiz_questions() — general/topic quiz via gpt-4o-mini JSON mode",
            "generate_quiz_from_chapter_content() — from textbook chapter text",
            "generate_quiz_from_chat() — from chat transcript + topic_summary",
            "Question schema: question, options[4], correct, explanation",
            "Validation: exactly 4 options, correct ∈ options",
        ],
    )
    doc.add_heading("11.2 Grade Subjects", 2)
    table(
        doc,
        ["Grade", "Subjects (API)"],
        [
            ["4", "Maths, General Science"],
            ["5", "Maths, General Science"],
            ["6", "Maths, General Science, Computer"],
            ["7", "Maths, General Science, Computer"],
        ],
    )
    doc.add_paragraph(
        "Note: quiz_engine uses 'Computer Science' for grades 6-8; API/RAG use 'Computer'. "
        "Chapter lookup for Computer requires matching BOOK_MAP key."
    )

    # 12. Analytics
    doc.add_heading("12. Analytics", 1)
    doc.add_heading("12.1 utils/quiz_db.get_user_analytics()", 2)
    bullets(
        doc,
        [
            "total_attempts, total_questions, total_correct, overall_accuracy",
            "total_time_minutes, avg_time_per_question, streak_days",
            "recent_attempts, subject_breakdown, daily_activity (14 days), performance_trend",
        ],
    )
    doc.add_heading("12.2 Parent Dashboard Computed Metrics", 2)
    bullets(
        doc,
        [
            "favourite_subject (most chat sessions)",
            "score_trend, speed_analysis (avg sec/question by subject)",
            "consistency = max(0, 100 - std_dev(scores))",
            "improvement = recent half avg − older half avg (if ≥4 attempts)",
        ],
    )
    doc.add_heading("12.3 Parent AI Report", 2)
    bullets(
        doc,
        [
            "GET /api/parent/report — GPT-4o-mini JSON mode (~1200 tokens)",
            "Sections: overview, strengths, improve, tips, next goals",
            "Fields: summary_headline, overall_rating, sections[]",
        ],
    )

    # 13. Speech
    doc.add_heading("13. Speech / TTS", 1)
    table(
        doc,
        ["Property", "Value"],
        [
            ["Endpoint", "POST /api/chat/speech"],
            ["Implementation", "chat_engine → utils/llm.text_to_speech_base64()"],
            ["Format", "MP3 base64"],
            ["Max text", "4000 chars (truncated)"],
            ["Storage", "In-memory only via this endpoint"],
        ],
    )

    # 14. Books
    doc.add_heading("14. Curriculum Books (books_mds)", 1)
    doc.add_paragraph(
        "Location: AutiStudy-React/books_mds/ (resolved by utils/book_parser.BOOKS_DIR). "
        "Used for quiz chapter selection; RAG uses separate ChromaDB index."
    )
    table(
        doc,
        ["Grade", "Subject", "File"],
        [
            ["4", "Maths", "Grade 4/math/math_4_parse.md"],
            ["4", "General Science", "Grade 4/gs/gs_4.md"],
            ["5", "Maths", "Grade 5/math/math_5.md"],
            ["5", "General Science", "Grade 5/gs/gs_5.md"],
            ["6", "Maths", "Grade 6/math/math_6.md"],
            ["6", "General Science", "Grade 6/gs/gs_6.md"],
            ["6", "Computer Science", "Grade 6/comp/comp_6.md"],
            ["7", "General Science", "Grade 7/gs/gs_7.md"],
            ["7", "Computer Science", "Grade 7/comp/comp_7.md"],
        ],
    )
    doc.add_paragraph(
        "Grade 7 Maths (math_parse_7.md) exists on disk but is NOT in BOOK_MAP."
    )
    doc.add_heading("14.1 Chapter Parsing (utils/book_parser.py)", 2)
    bullets(
        doc,
        [
            "Heading patterns: # Chapter 01, # Unit 1, # Sub-Domain, ## Q1, etc.",
            "Filter hits with start_line >= 100 (skip TOC)",
            "get_chapter_content() returns up to 6000 chars, strips HTML tables",
        ],
    )

    # 15. Teaching Agent & Emotion
    doc.add_heading("15. Teaching Agent & Emotion (Pre–Media Agent)", 1)
    doc.add_heading("15.1 teaching_agent.py — Rule-Based Modality", 2)
    bullets(
        doc,
        [
            "Modality ladder: text → text_image → text_image_voice → step_by_step",
            "≥2 consecutive confused/frustrated → escalate one level",
            "≥3 consecutive happy/focused → de-escalate one level",
            "neutral/no_face → counters unchanged",
            "In-memory state keyed by session_id",
        ],
    )
    doc.add_heading("15.2 emotion.py — Vision Service", 2)
    bullets(
        doc,
        [
            "Model: gpt-4o-mini Vision, detail: low",
            "Input: base64 JPEG webcam frame",
            "Emotions: happy, focused, neutral, confused, frustrated, no_face",
            "Output: emotion, confidence, understood, description",
        ],
    )
    doc.add_heading("15.3 agent_memory.py", 2)
    bullets(
        doc,
        [
            "Cross-session JSON per student (MD5 of email)",
            "Tracks: tool_success/failure, struggling_topics, mastered_topics, modality_success",
            "recent_sessions (last 5), confusion resolution rate",
            "POST /api/agent/session-summary, GET /api/agent/memory",
        ],
    )

    # 16. OCR
    doc.add_heading("16. OCR Service (utils/ocr.py)", 1)
    bullets(
        doc,
        [
            "Engine: EasyOCR (en, CPU)",
            "Preprocessing: grayscale, upscale 1500px, adaptive threshold, dilate",
            "NADRA CRC/B-Form validation",
            "Parent CNIC extraction by frequency analysis",
        ],
    )

    # 17. Dependencies
    doc.add_heading("17. Dependencies (requirements.txt)", 1)
    table(
        doc,
        ["Package", "Purpose"],
        [
            ["fastapi, uvicorn, pydantic", "REST API"],
            ["openai", "LLM, TTS, Vision, images"],
            ["chromadb, sentence-transformers", "Vector store + embeddings"],
            ["rank_bm25, numpy", "Sparse retrieval + scores"],
            ["bcrypt", "Password hashing"],
            ["streamlit, plotly, pandas", "Legacy UI + charts"],
            ["Pillow", "Image handling"],
            ["replicate", "Legacy alt image (llm2.py)"],
        ],
    )
    doc.add_paragraph("Optional (not in requirements.txt): easyocr, opencv-python for B-Form OCR.")

    # 18. Environment
    doc.add_heading("18. Environment Variables", 1)
    table(
        doc,
        ["Variable", "Default", "Purpose"],
        [
            ["OPENAI_API_KEY", "—", "Required for LLM, TTS, Vision, images"],
            ["CHROMA_DB_PATH", "OneSharedChromaDB", "ChromaDB directory"],
            ["CHROMA_COLLECTION", "ptb_textbooks", "Collection name"],
            ["PYTHONIOENCODING", "utf-8", "Set by run_api_detached.py"],
        ],
    )
    doc.add_heading("18.1 Key Sources (priority)", 2)
    bullets(
        doc,
        [
            "1. OPENAI_API_KEY environment variable",
            "2. .streamlit/secrets.toml",
            "3. Plain files: new_imp_fyp_open_ai_key.txt, openai_api_key.txt, OPENAI_API_KEY.txt",
        ],
    )

    # 19. Running
    doc.add_heading("19. Running the Backend", 1)
    code_block(
        doc,
        "cd AutiStudy\n"
        "uvicorn api_server:app --port 8000 --reload\n\n"
        "# From React repo:\n"
        "npm run dev:api\n"
        "# Runs: uvicorn api_server:app --app-dir ../AutiStudy --port 8000",
    )
    doc.add_paragraph(
        "Windows detached: python scripts/run_api_detached.py (uses pythonw.exe)."
    )

    # 20. Streamlit Legacy
    doc.add_heading("20. Streamlit Legacy Application", 1)
    bullets(
        doc,
        [
            "app.py — multi-page Streamlit router",
            "views/: landing, login, signup, dashboard, ai_tutor, chat, analytics, practice_quiz, faq, about",
            "Same JSON files and utils modules as FastAPI",
            "Dockerfile deploys Streamlit on port 7860 (not FastAPI)",
        ],
    )

    # 21. Media Agent Exclusion
    doc.add_heading("21. Media Agent — Exclusion Boundary", 1)
    doc.add_paragraph(
        "The following is intentionally OUT OF SCOPE for this document:"
    )
    table(
        doc,
        ["Item", "Location"],
        [
            ["Media Agent module", "utils/media_agent.py (~590 lines)"],
            ["ReAct API endpoint", "POST /api/agent/run"],
            ["Pattern", "GPT-4o function calling, max 3 iterations"],
            ["Tools", "do_nothing, simplify_text, generate_visual, speak_aloud, explain_steps, use_analogy, check_prerequisite, notify_parent"],
        ],
    )
    doc.add_paragraph(
        "Documented above (NOT media agent): teaching_agent.py, emotion.py, agent_memory.py, "
        "chat_engine, visual_aids, POST /api/agent/analyze-emotion, POST /api/chat/sessions/{id}/image."
    )

    # 22. Known Issues
    doc.add_heading("22. Known Implementation Notes", 1)
    bullets(
        doc,
        [
            "Dual frontend sync: FastAPI and Streamlit share JSON — no ORM",
            "Subject naming inconsistency: 'Computer' vs 'Computer Science'",
            "RAT disabled in production API for latency",
            "Grade 7 Maths book not in BOOK_MAP",
            "Debug endpoint /api/debug/openai-ping should be removed before production",
            "Dockerfile deploys Streamlit, not FastAPI — React expects API on port 8000",
        ],
    )

    return doc


def build_frontend() -> Document:
    doc = Document()
    style_doc(doc)
    title_page(
        doc,
        "AutiStudy Frontend",
        "Technical documentation — Next.js React web application (autistudy-web)",
    )

    # 1. Overview
    doc.add_heading("1. Executive Summary", 1)
    doc.add_paragraph(
        "AutiStudy-React (package name autistudy-web v0.1.0) is a Next.js 14 App Router "
        "frontend for an AI-powered adaptive learning platform targeting autistic students "
        "in Pakistan (grades 4–7). It communicates with a FastAPI backend at http://127.0.0.1:8000 "
        "via a typed fetch client. Features include bilingual EN/Urdu UI, RAG-grounded chat with "
        "11 in-browser visual aid types, quizzes, analytics, parent dashboards, webcam teaching "
        "agent integration, and autism-friendly accessibility settings."
    )

    # 2. Stack
    doc.add_heading("2. Technology Stack", 1)
    table(
        doc,
        ["Layer", "Technology", "Version"],
        [
            ["Framework", "Next.js (App Router)", "14.2.35"],
            ["UI", "React", "18"],
            ["Language", "TypeScript", "5"],
            ["Styling", "Tailwind CSS", "3.4.1"],
            ["Animation", "Framer Motion", "12.38"],
            ["Scroll", "Lenis", "1.3"],
            ["Landing anim", "GSAP", "3.15"],
            ["Markdown/Math", "react-markdown, KaTeX", "—"],
            ["Icons", "lucide-react", "1.8"],
            ["E2E", "Playwright", "1.59 (dev)"],
            ["Backend", "FastAPI sibling AutiStudy/", "port 8000"],
        ],
    )

    # 3. Structure
    doc.add_heading("3. Project Structure", 1)
    code_block(
        doc,
        "AutiStudy-React/\n"
        "├── app/                    # Next.js App Router pages\n"
        "│   ├── layout.tsx          # Root layout + providers\n"
        "│   ├── page.tsx            # Landing (/)\n"
        "│   ├── login/, signup/, dashboard/, chat/, quiz/, analytics/\n"
        "│   └── parent/dashboard/, parent/report/\n"
        "├── components/\n"
        "│   ├── agent/              # EmotionCamera, CameraConsent, UnderstandingCheck\n"
        "│   ├── landing/            # HeroIntro, Features, HowItWorks, CTABanner\n"
        "│   ├── layout/             # NavBar, Footer, PageShell\n"
        "│   ├── primitives/         # Aurora, DancingButton, SmoothScrollProvider\n"
        "│   └── settings/           # SettingsModal\n"
        "├── lib/\n"
        "│   ├── api/client.ts       # Typed API client\n"
        "│   ├── auth/               # AuthProvider, redirect helpers\n"
        "│   ├── hooks/useTeachingAgent.ts\n"
        "│   ├── i18n/               # LocaleProvider, translations\n"
        "│   └── settings/SettingsContext.tsx\n"
        "├── scripts/                # Playwright verification scripts\n"
        "├── books_mds/              # Curriculum markdown (backend RAG source)\n"
        "├── next.config.mjs, tailwind.config.ts, package.json\n",
    )
    doc.add_paragraph("Path alias: @/* maps to project root (tsconfig.json).")

    # 4. Routes
    doc.add_heading("4. Application Routes", 1)
    table(
        doc,
        ["Route", "File", "Auth", "Purpose"],
        [
            ["/", "app/page.tsx", "Public", "Landing page"],
            ["/about", "app/about/page.tsx", "Public", "About page"],
            ["/faq", "app/faq/page.tsx", "Public", "FAQ accordion"],
            ["/login", "app/login/page.tsx", "Public", "Student/parent login"],
            ["/signup", "app/signup/page.tsx", "Public", "Student/parent signup"],
            ["/dashboard", "app/dashboard/page.tsx", "Student", "Stats, subjects, resume chats"],
            ["/chat", "app/chat/page.tsx", "Student", "AI tutor (~2600 lines)"],
            ["/quiz", "app/quiz/page.tsx", "Student", "Practice quiz flow"],
            ["/analytics", "app/analytics/page.tsx", "Student", "Progress charts"],
            ["/parent/dashboard", "app/parent/dashboard/page.tsx", "Parent token", "Child metrics"],
            ["/parent/report", "app/parent/report/page.tsx", "Parent token", "AI progress report"],
        ],
    )
    doc.add_paragraph(
        "No middleware.ts — all auth guards are client-side useEffect redirects."
    )

    # 5. Layout & Providers
    doc.add_heading("5. Root Layout & Providers", 1)
    doc.add_paragraph("File: app/layout.tsx")
    bullets(
        doc,
        [
            "Fonts: Quicksand (display), Inter (body) via next/font/google",
            "Provider nesting: LocaleProvider → SettingsProvider → AuthProvider → SmoothScrollProvider",
            "Global: Aurora background, SettingsModal overlay",
            "Metadata: title, description, favicon",
        ],
    )

    # 6. Auth
    doc.add_heading("6. Authentication", 1)
    doc.add_heading("6.1 AuthProvider (lib/auth/AuthProvider.tsx)", 2)
    bullets(
        doc,
        [
            "Token: localStorage key autistudy_token",
            "Mount: if token → authApi.me(); 401 clears token silently",
            "Methods: login, register, logout, refresh",
            "Hook: useAuth() — throws outside provider",
        ],
    )
    doc.add_heading("6.2 Redirect Helpers (lib/auth/redirect.ts)", 2)
    bullets(
        doc,
        [
            "loginUrlFor(currentPath?) → /login?next=<encoded path>",
            "safeNext(raw) — prevents open redirects (only same-origin paths starting with /)",
            "Preserves deep links e.g. /chat?session=abc across login",
        ],
    )
    doc.add_heading("6.3 Dual Token Model", 2)
    table(
        doc,
        ["Role", "Storage Key", "API Helper"],
        [
            ["Student", "autistudy_token", "api() with Bearer header"],
            ["Parent", "autistudy_parent_token", "apiParent()"],
        ],
    )
    doc.add_heading("6.4 Login Flow", 2)
    bullets(
        doc,
        [
            "Role picker: Student vs Parent",
            "Student: useAuth().login() → redirect to ?next= or /dashboard",
            "Parent: parentApi.login() → setParentToken → /parent/dashboard (ignores ?next=)",
        ],
    )
    doc.add_heading("6.5 Signup Flow", 2)
    bullets(
        doc,
        [
            "Student: ChildSignupForm — multipart B-Form via parentApi.childSignup()",
            "Fields: name, email, password, grade 4-7, CNIC, B-Form photo",
            "Parent: ParentSignupForm — parentApi.signup() with child CNIC verification",
        ],
    )

    # 7. API Client
    doc.add_heading("7. API Client (lib/api/client.ts)", 1)
    doc.add_paragraph("API_BASE = process.env.NEXT_PUBLIC_API_URL || http://127.0.0.1:8000")
    doc.add_heading("7.1 Core api() Function", 2)
    bullets(
        doc,
        [
            "JSON Content-Type, optional Authorization Bearer",
            "Parses FastAPI detail (string or Pydantic validation array)",
            "Network failure → ApiError(0, friendly message)",
            "Handles 204 empty responses",
        ],
    )
    doc.add_heading("7.2 API Groups", 2)
    table(
        doc,
        ["Group", "Endpoints"],
        [
            ["authApi", "/api/auth/register, /login, /logout, /me"],
            ["userApi", "/api/users/me/stats, /subjects, /recent-chats, /recent-quizzes"],
            ["chatApi", "/api/chat/config, /sessions, messages, image, speech, quiz"],
            ["quizApi", "/api/quiz/chapters, generate, generate-from-chat, submit, /analytics"],
            ["parentApi", "/api/auth/parent/*, /api/parent/dashboard, /report, child signup"],
        ],
    )
    doc.add_heading("7.3 Visual Aid Types (TypeScript interfaces)", 2)
    bullets(
        doc,
        [
            "MathStepCard, EmojiCountingData, FactorTreeData, FractionBarData",
            "NumberLineData, BarChartData, PercentageBarData, TimesTableData",
            "GeometryData, RatioData",
            "GenerateVisualAidResponse — discriminated union by kind",
            "resolveImageUrl() — prefixes relative /api/generated-images/ paths",
        ],
    )

    # 8. Chat Page
    doc.add_heading("8. Chat Page (app/chat/page.tsx)", 1)
    doc.add_heading("8.1 URL State Machine", 2)
    table(
        doc,
        ["URL", "State", "Behavior"],
        [
            ["/chat", "Picker", "Subject selection from userApi.subjects()"],
            ["/chat?subject=Maths", "Bootstrap", "POST create session → redirect ?session=id"],
            ["/chat?session=id", "Conversation", "Load session, send messages, multimedia"],
        ],
    )
    doc.add_heading("8.2 Conversation Features", 2)
    bullets(
        doc,
        [
            "Optimistic UI for user messages with rollback on error",
            "Markdown + KaTeX via react-markdown, remark-math, rehype-katex",
            "normalizeMath() for \\(...\\) and \\[...\\] delimiters",
            "11 in-browser visual views: EmojiCounting, MathStepCard, FactorTree, FractionBar, NumberLine, BarChart, PercentageBar, TimesTable, Geometry, Ratio + DALL·E images",
            "TTS: chatApi.speak() → base64 MP3 in audio element",
            "Action buttons only on last assistant message (cost control)",
            "ChatQuizModal when ≥4 messages",
            "Config banner when tutor_configured === false",
            "Language for new sessions: locale ur → ur else en",
        ],
    )
    doc.add_heading("8.3 Layout", 2)
    bullets(
        doc,
        [
            "ChatShell: NavBar + pr-72 for fixed EmotionCamera panel",
            "Fixed EmotionCamera on right edge",
        ],
    )

    # 9. Dashboard
    doc.add_heading("9. Dashboard (app/dashboard/page.tsx)", 1)
    bullets(
        doc,
        [
            "Parallel fetch: stats, subjects, recentChats, recentQuizzes",
            "Animated header with avatar, word-stagger greeting, star badge (useCountUp)",
            "4 stat cards: stars, streak, quizzes, accuracy",
            "Quick actions: Quiz, Analytics",
            "Subject cards → /chat?subject=<name>",
            "Continue learning (top 3 chats) → /chat?session=<id>",
            "formatRelative() via Intl.RelativeTimeFormat (en/ur)",
        ],
    )

    # 10. Quiz
    doc.add_heading("10. Quiz Page (app/quiz/page.tsx)", 1)
    bullets(
        doc,
        [
            "Phases: pick → chapters → loading → question → result",
            "Subjects by grade: 4-5 Maths+GS; 6-8 +Computer Science",
            "Sources: generateFromChat OR by chapter via quizApi.chapters()",
            "MCQ reveal → explanation → timed per-question → submit",
            "Result: animated score ring, star rating, question review",
        ],
    )

    # 11. Analytics
    doc.add_heading("11. Analytics (app/analytics/page.tsx)", 1)
    bullets(
        doc,
        [
            "Data: quizApi.analytics() → AnalyticsData",
            "StarsBanner, stat grid, ActivityChart (7 days), SubjectBreakdown",
            "RecentAttempts (last 8), empty state CTA to /quiz",
        ],
    )

    # 12. Parent Pages
    doc.add_heading("12. Parent Pages", 1)
    doc.add_heading("12.1 Parent Dashboard", 2)
    bullets(
        doc,
        [
            "Auth: getParentToken(); 401 → clear → /login",
            "Custom SVG: DonutChart, ScoreTrend, SpeedBar",
            "Subject accuracy bars, daily activity, expandable quiz history",
            "Link to /parent/report",
        ],
    )
    doc.add_heading("12.2 Parent Report", 2)
    bullets(
        doc,
        [
            "parentApi.report() → parsed ReportData sections",
            "Overall rating: Excellent / Good / Developing / Needs Support",
            "SectionCard per report section with colored themes",
            "i18n via t.parentReport",
        ],
    )

    # 13. Teaching Agent
    doc.add_heading("13. Teaching Agent Integration", 1)
    doc.add_heading("13.1 useTeachingAgent (lib/hooks/useTeachingAgent.ts)", 2)
    bullets(
        doc,
        [
            "Modalities: text → text_image → text_image_voice → step_by_step",
            "getUserMedia 320×240, capture JPEG every 4 seconds",
            "POST /api/agent/run with session_id, image_b64, consecutive_confused, tools_used",
            "Camera consent: localStorage autistudy_camera_consent (granted|denied|pending)",
            "Exports: state, videoRef, startCamera, stopCamera, forceModality",
        ],
    )
    doc.add_heading("13.2 Agent UI Components", 2)
    table(
        doc,
        ["Component", "File", "Purpose"],
        [
            ["EmotionCamera", "components/agent/EmotionCamera.tsx", "Fixed right panel: video, emotion badge, modality, plan"],
            ["CameraConsentModal", "components/agent/CameraConsentModal.tsx", "One-time consent dialog (1.5s delay)"],
            ["UnderstandingCheck", "components/agent/UnderstandingCheck.tsx", "Yes/Not yet after each assistant message"],
        ],
    )
    doc.add_heading("13.3 Agent-Triggered Chat Actions", 2)
    bullets(
        doc,
        [
            "generate_visual → auto onGenerateImage()",
            "speak_aloud → auto onSpeak() TTS",
            "explain_steps, simplify_text, use_analogy, check_prerequisite → inject follow-up via onSend()",
        ],
    )

    # 14. i18n
    doc.add_heading("14. Internationalization", 1)
    bullets(
        doc,
        [
            "Locales: en | ur (lib/i18n/LocaleProvider.tsx)",
            "Storage: localStorage autistudy_locale",
            "DOM: document.documentElement.lang and .dir (rtl for Urdu)",
            "translations.ts: ~508 lines mirrored en/ur objects",
            "Namespaces: brand, nav, intro, hero, features, auth, pages.*, parentDashboard, parentReport",
            "Urdu font: Noto Nastaliq Urdu in globals.css",
            "Toggle: NavBar globe button + Settings modal",
        ],
    )

    # 15. Settings
    doc.add_heading("15. Settings System", 1)
    table(
        doc,
        ["Setting", "Type", "Effect"],
        [
            ["fontSize", "normal|large|xl", "html.font-large / font-xl (18px/21px)"],
            ["reduceMotion", "boolean", "html.reduce-motion"],
            ["highContrast", "boolean", "html.high-contrast"],
            ["focusMode", "boolean", "html.focus-mode — hides aurora"],
            ["ttsAutoRead", "boolean", "Stored only — not yet wired to chat auto-TTS"],
        ],
    )
    doc.add_paragraph("SettingsModal: global overlay z-200 with Appearance, Accessibility, Language, Profile, Account (password change), About sections.")

    # 16. UI Components
    doc.add_heading("16. UI Components", 1)
    doc.add_heading("16.1 Layout", 2)
    bullets(
        doc,
        [
            "NavBar: fixed glass nav, auth-aware links, 5s delay on / only, locale toggle, settings, logout",
            "Footer: tagline + copyright",
            "PageShell: standard wrapper for About/FAQ",
        ],
    )
    doc.add_heading("16.2 Landing", 2)
    bullets(
        doc,
        [
            "HeroIntro: 3-phase intro animation (big → logo → hero), GSAP glitter letters",
            "Features, HowItWorks, CTABanner",
            "Hero visual variants: EducationalVisual, KidLearningVisual, PhotoVisual",
            "Unsplash images via next.config.mjs remotePatterns",
        ],
    )
    doc.add_heading("16.3 Primitives", 2)
    bullets(
        doc,
        [
            "Aurora: animated gradient + floating blur orbs",
            "SmoothScrollProvider: Lenis, exposes window.__lenis",
            "DancingButton: primary/ghost/soft CTA with hover wiggle",
        ],
    )

    # 17. Styling
    doc.add_heading("17. Styling", 1)
    doc.add_heading("17.1 Tailwind (tailwind.config.ts)", 2)
    bullets(
        doc,
        [
            "Palette: glacier (50-700), mint, deep text colors",
            "font-display (Quicksand), font-sans (Inter)",
            "Shadows: glow, soft, deep",
            "Animations: drift-bg, float-slow, breathe, bounce-soft, shimmer",
        ],
    )
    doc.add_heading("17.2 globals.css", 2)
    bullets(
        doc,
        [
            "KaTeX import, body #F7FBFC / #0F2D4A",
            "Lenis scroll overrides, RTL Urdu font",
            "Accessibility: .font-large, .font-xl, .reduce-motion, .high-contrast, .focus-mode",
            "Glass utilities: .glass, .glass-strong, .text-shimmer",
            "Landing glitter: .glitter-letter, .sparkle, rainbow-shift keyframes",
            "prefers-reduced-motion media query",
        ],
    )
    doc.add_heading("17.3 Design Language", 2)
    doc.add_paragraph(
        "Autism-friendly: calm glacier/mint palette, glass morphism, predictable motion, "
        "generous spacing, large touch targets on understanding checks and quiz options."
    )

    # 18. Next Config
    doc.add_heading("18. Configuration", 1)
    bullets(
        doc,
        [
            "next.config.mjs: images.remotePatterns for images.unsplash.com",
            "No API rewrites — frontend calls backend directly (CORS on FastAPI)",
            "postcss.config.mjs: tailwindcss only",
            ".eslintrc.json: next/core-web-vitals, next/typescript",
        ],
    )

    # 19. NPM Scripts
    doc.add_heading("19. NPM Scripts", 1)
    table(
        doc,
        ["Script", "Command"],
        [
            ["dev", "next dev"],
            ["dev:api", "uvicorn api_server:app --app-dir ../AutiStudy --port 8000"],
            ["dev:all", "concurrently web + api"],
            ["build", "next build"],
            ["start", "next start"],
            ["lint", "next lint"],
        ],
    )

    # 20. Verification Scripts
    doc.add_heading("20. Playwright Verification Scripts", 1)
    doc.add_paragraph(
        "Located in scripts/. Run with node scripts/<name>.mjs. "
        "Defaults: WEB http://localhost:3000, API http://127.0.0.1:8000."
    )
    table(
        doc,
        ["Script", "Purpose"],
        [
            ["verify-auth-redirect.mjs", "Login preserves ?next= deep links"],
            ["verify-dashboard-e2e.mjs", "Signup → dashboard → logout"],
            ["verify-chat-e2e.mjs", "Dashboard → chat → message → resume"],
            ["verify-logout.mjs", "Logout clears token"],
            ["verify-chat-picker.mjs", "Subject picker flow"],
            ["verify-chat-real-llm.mjs", "Real LLM + config banner"],
            ["verify-chat-rag-multimedia.mjs", "RAG flags, read-aloud, picture buttons"],
            ["verify-visual-aid-router.mjs", "Image vs step-card routing"],
            ["verify-division-method.mjs", "Long division step cards"],
            ["verify-math-render.mjs", "KaTeX rendering"],
            ["verify-markdown-render.mjs", "Markdown with stubbed API"],
            ["verify-button-row.mjs", "Visual aid button row"],
            ["verify-intro-timing.mjs", "Hero intro phase durations"],
            ["capture-*.mjs, snap-*.mjs", "Design screenshot capture"],
        ],
    )
    doc.add_paragraph("Output: design-captures/, design-captures-variants/")

    # 21. Environment
    doc.add_heading("21. Environment Variables", 1)
    table(
        doc,
        ["Variable", "Default", "Purpose"],
        [
            ["NEXT_PUBLIC_API_URL", "http://127.0.0.1:8000", "Backend base URL"],
            ["WEB_BASE_URL", "http://localhost:3000", "Scripts only"],
            ["API_BASE_URL", "http://127.0.0.1:8000", "Scripts only"],
        ],
    )

    # 22. Data Flow
    doc.add_heading("22. Data Flows", 1)
    doc.add_heading("22.1 Student Chat Flow", 2)
    bullets(
        doc,
        [
            "User visits /chat?session=id",
            "getToken() from localStorage",
            "If no token → loginUrlFor() redirect",
            "GET /api/chat/sessions/id + GET /api/chat/config",
            "POST /api/chat/sessions/id/messages on send",
            "POST /api/chat/sessions/id/image for visual aids",
            "POST /api/chat/speech for TTS",
        ],
    )
    doc.add_heading("22.2 Teaching Agent Loop", 2)
    bullets(
        doc,
        [
            "getUserMedia → capture frame every 4s",
            "POST /api/agent/run",
            "Frontend applies tool side effects (generate_visual, speak_aloud, onSend follow-ups)",
            "UnderstandingCheck after each assistant message",
        ],
    )

    # 23. Gaps & Notes
    doc.add_heading("23. Implementation Notes & Gaps", 1)
    bullets(
        doc,
        [
            "No Next.js middleware — client-side guards only",
            "Chat page is a monolith (~2674 lines) — all visual views in one file",
            "ttsAutoRead setting not yet consumed in chat",
            "Student signup uses B-Form multipart, not AuthProvider.register()",
            "Parent login ignores ?next= return URL",
            "Quiz page defines grade 8 subjects but signup only offers 4-7",
            "books_mds/ is backend RAG source, not imported by frontend",
        ],
    )

    return doc


def main() -> None:
    backend_path = OUT_DIR / "backend.docx"
    frontend_path = OUT_DIR / "frontend.docx"

    build_backend().save(str(backend_path))
    build_frontend().save(str(frontend_path))

    print(f"Created: {backend_path}")
    print(f"Created: {frontend_path}")


if __name__ == "__main__":
    main()
